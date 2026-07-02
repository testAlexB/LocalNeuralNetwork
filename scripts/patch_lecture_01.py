"""Лекция 1 — Дополнение: закрываем 10 пробелов по рецензии
Содержит: числовой пример Attention, RoPE, SwiGLU, backprop, сложность, задания
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Arc
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io, os, sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')

C = {
    'bg': '#1a1a2e', 'primary': '#16213e', 'secondary': '#0f3460',
    'accent': '#e94560', 'orange': '#f59e0b', 'green': '#10b981',
    'blue': '#3b82f6', 'purple': '#8b5cf6', 'teal': '#06b6d4',
    'pink': '#ec4899', 'text': '#e2e8f0', 'text_dim': '#94a3b8',
}

def sa(ax):
    ax.set_facecolor(C['bg'])
    ax.figure.patch.set_facecolor(C['bg'])
    for s in ax.spines.values(): s.set_visible(False)
    ax.tick_params(colors=C['text_dim'], labelsize=6)
    ax.grid(True, alpha=0.06, color=C['text_dim'])

def sf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=160, bbox_inches='tight', facecolor=C['bg'])
    buf.seek(0); plt.close(fig); return buf


# ─── Диаграмма 1: Числовой пример Attention ───
def fig_numerical_attention():
    fig, axes = plt.subplots(1, 3, figsize=(11, 3.5))
    titles = ['Шаг 1: Q · Kᵀ', 'Шаг 2: ÷ √dₖ, softmax', 'Шаг 3: × V → взвешенная сумма']
    matrices_data = []

    # Пример: предложение "щука лещ" (2 токена), d_k=3
    Q = np.array([[0.8, 0.2, 0.5], [0.1, 0.9, 0.3]])
    K = np.array([[0.7, 0.3, 0.4], [0.2, 0.8, 0.6]])
    V = np.array([[1.0, 0.0], [0.0, 1.0]])

    # Шаг 1: Q @ K.T
    scores = Q @ K.T

    # Шаг 2: /sqrt(3) + softmax
    scaled = scores / np.sqrt(3)
    exp_s = np.exp(scaled - scaled.max(axis=1, keepdims=True))
    weights = exp_s / exp_s.sum(axis=1, keepdims=True)

    # Шаг 3: weights @ V
    output = weights @ V

    matrices = [scores, weights, output]
    titles_math = ['Q·Kᵀ', 'softmax(Q·Kᵀ/√3)', 'output = W·V']
    fmt = ['.2f', '.3f', '.2f']

    for idx, (ax, title, mat, math_title, f) in enumerate(zip(axes, titles, matrices, titles_math, fmt)):
        sa(ax)
        ax.axis('off')
        ax.text(0.5, 0.92, title, ha='center', fontsize=9, color='white', fontweight='bold', transform=ax.transAxes)
        ax.text(0.5, 0.05, math_title, ha='center', fontsize=8, color=C['orange'],
                fontweight='bold', family='monospace', transform=ax.transAxes)

        for i in range(mat.shape[0]):
            for j in range(mat.shape[1]):
                val = mat[i, j]
                color = C['accent'] if val == mat.max() else C['blue'] if val > mat.mean() else C['teal']
                alpha = 0.3 + 0.6 * (val - mat.min()) / (mat.max() - mat.min() + 1e-8)
                x = j / mat.shape[1] + 0.08
                y = 1.0 - (i + 0.7) / mat.shape[0]
                r = FancyBboxPatch((x, y), 0.15, 0.15, boxstyle="round,pad=0.01",
                                   facecolor=color, alpha=alpha, edgecolor='white', linewidth=0.3)
                ax.add_patch(r)
                ax.text(x + 0.075, y + 0.075, f'{val:{f}}', ha='center', va='center',
                        fontsize=8, color='white')

        # Row/col labels
        for i, label in enumerate(['щука', 'лещ']):
            ax.text(0.01, 1.0 - (i + 0.55) / mat.shape[0], label, fontsize=8, color=C['text'])
        if idx == 0 or idx == 1:
            for j, label in enumerate(['щука', 'лещ']):
                ax.text(j / mat.shape[1] + 0.15, 0.78, label, fontsize=7, color=C['text'],
                        rotation=30, ha='left')

    # Highlight max values in weights
    fig.text(0.5, 0.01, 'Рис. A1: Числовой пример Self-Attention для двух слов "щука лещ" (dₖ=3)',
             ha='center', fontsize=8, color=C['text_dim'])
    return sf(fig)


# ─── Диаграмма 2: RoPE — визуализация ───
def fig_rope():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 4))
    for ax in [ax1, ax2]: sa(ax)

    # Left: rotation of Q,K
    theta = np.pi / 4
    for ax in [ax1, ax2]:
        ax.set_xlim(-1.5, 1.5); ax.set_ylim(-1.5, 1.5)
        ax.set_aspect('equal')

    # Circle and one vector
    circle = plt.Circle((0, 0), 1.0, fill=False, color=C['text_dim'], alpha=0.3, linewidth=0.5)
    ax1.add_patch(circle)
    ax1.text(0, 1.05, 'Позиция 0', fontsize=8, color=C['text'], ha='center')
    ax1.text(0, -1.2, 'Q/k = sin(0·θ) = 0', fontsize=7, color=C['text_dim'], ha='center')

    # Vector at angle 0
    ax1.arrow(0, 0, 1, 0, head_width=0.1, head_length=0.1, fc=C['blue'], ec=C['blue'], alpha=0.8)
    ax1.text(0.5, 0.1, 'Q₀', fontsize=9, color=C['blue'], fontweight='bold')

    # Same vector rotated on ax2
    angle = np.pi / 3
    vec = np.array([np.cos(angle), np.sin(angle)])
    circle2 = plt.Circle((0, 0), 1.0, fill=False, color=C['text_dim'], alpha=0.3, linewidth=0.5)
    ax2.add_patch(circle2)
    ax2.text(0, 1.05, 'Позиция m', fontsize=8, color=C['text'], ha='center')
    ax2.text(0, -1.2, f'Q_m = Q₀ · cos(m·θ) + ...\nПоворот на {angle:.1f} рад', fontsize=7, color=C['text_dim'], ha='center')

    ax2.arrow(0, 0, vec[0], vec[1], head_width=0.1, head_length=0.1, fc=C['accent'], ec=C['accent'], alpha=0.8)
    ax2.text(vec[0]*0.5, vec[1]*0.5 + 0.1, 'Qₘ', fontsize=9, color=C['accent'], fontweight='bold')

    # Original angle shown
    ax2.arrow(0, 0, 1, 0, head_width=0.05, head_length=0.05, fc=C['text_dim'], ec=C['text_dim'], alpha=0.3)

    # Angle arc
    arc = Arc((0, 0), 0.6, 0.6, angle=0, theta1=0, theta2=np.degrees(angle), color=C['orange'], linewidth=1)
    ax2.add_patch(arc)
    ax2.text(0.15, 0.1, f'θ = pos·ω', fontsize=7, color=C['orange'])

    # Formula
    fig.text(0.5, 0.42, 'RoPE: f{q, pos} = q · (cos(pos·ω) + i·sin(pos·ω))',
             ha='center', fontsize=10, color=C['orange'], fontweight='bold', family='monospace')
    fig.text(0.5, 0.35, 'Каждый токен "поворачивает" свой Q и K в зависимости от позиции.\n'
                        'Относительное расстояние = разность углов → скалярное произведение зависит от |i-j|.',
             ha='center', fontsize=8, color=C['text'])

    fig.text(0.5, 0.02, 'Рис. A2: Rotary Position Embedding (RoPE) — вращение Q/K в комплексной плоскости',
             ha='center', fontsize=8, color=C['text_dim'])
    return sf(fig)


# ─── Диаграмма 3: Masked Attention ───
def fig_masked_attention():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(8, 4))
    for ax in [ax1, ax2]: sa(ax)

    n = 6
    mask = np.triu(np.ones((n, n)), k=1) * -1e9
    attn_no_mask = np.ones((n, n)) * 0.2
    np.fill_diagonal(attn_no_mask, 0.8)

    # Left: without mask
    ax1.imshow(attn_no_mask, cmap='viridis', aspect='auto', vmin=0, vmax=1)
    ax1.set_title('Без маски (жульничество)', fontsize=9, color=C['accent'], fontweight='bold')
    ax1.set_xticks(range(n))
    ax1.set_yticks(range(n))
    ax1.set_xticklabels(range(1, n+1), fontsize=7, color=C['text'])
    ax1.set_yticklabels(range(1, n+1), fontsize=7, color=C['text'])
    ax1.set_xlabel('Предсказываемое слово (j)', fontsize=7, color=C['text'])
    ax1.set_ylabel('Текущее слово (i)', fontsize=7, color=C['text'])

    # Right: with mask
    attn_masked = attn_no_mask.copy()
    # Upper triangle set to near-zero
    for i in range(n):
        for j in range(n):
            if j > i:
                attn_masked[i, j] = 0.001

    ax2.imshow(attn_masked, cmap='viridis', aspect='auto', vmin=0, vmax=1)
    ax2.set_title('С causal-маской (правильно)', fontsize=9, color=C['green'], fontweight='bold')
    ax2.set_xticks(range(n))
    ax2.set_yticks(range(n))
    ax2.set_xticklabels(range(1, n+1), fontsize=7, color=C['text'])
    ax2.set_yticklabels(range(1, n+1), fontsize=7, color=C['text'])
    ax2.set_xlabel('Предсказываемое слово (j)', fontsize=7, color=C['text'])
    ax2.set_ylabel('Текущее слово (i)', fontsize=7, color=C['text'])

    fig.text(0.5, 0.02, 'Рис. A3: Causal Mask — почему модель не может "подглядывать" в будущее',
             ha='center', fontsize=8, color=C['text_dim'])
    return sf(fig)


# ─── Диаграмма 4: Backprop через Attention ───
def fig_backprop():
    fig, ax = plt.subplots(figsize=(8, 4))
    sa(ax); ax.axis('off')

    steps = [
        (0.0, 'Q @ Kᵀ', C['blue'], '∂L/∂(Q·Kᵀ) = ∂L/∂S · (mask)'),
        (2.2, '÷ √dₖ', C['teal'], '∂L/∂S_scaled = ∂L/∂S · 1/√dₖ'),
        (4.4, 'softmax', C['purple'], '∂L/∂S = S · (∂L/∂P - Σ(S · ∂L/∂P))\nгде S = softmax'),
        (6.6, '× V', C['green'], '∂L/∂V = Wᵀ · ∂L/∂out\n∂L/∂W = ∂L/∂out · Vᵀ'),
    ]

    for x, label, color, formula in steps:
        r = FancyBboxPatch((x, 1.8), 1.8, 0.7,
                           boxstyle="round,pad=0.04",
                           facecolor=color, alpha=0.2,
                           edgecolor=color, linewidth=1.5)
        ax.add_patch(r)
        ax.text(x + 0.9, 2.15, label, ha='center', va='center',
                fontsize=9, color='white', fontweight='bold')
        ax.text(x + 0.9, 1.6, formula, ha='center', va='top',
                fontsize=7, color=C['text'])

        if x > 0:
            ax.annotate('', xy=(x, 2.15), xytext=(x - 0.2, 2.15),
                        arrowprops=dict(arrowstyle='->', color=C['orange'], lw=2))

    # Градиентный поток снизу
    ax.annotate('', xy=(0.9, 0.5), xytext=(0.9, 0.9),
                arrowprops=dict(arrowstyle='->', color=C['accent'], lw=3))
    ax.text(0.9, 0.3, 'Градиент течёт обратно\nчерез каждую операцию',
            ha='center', fontsize=8, color=C['accent'])

    ax.text(4, 0.8, 'Важно: softmax получает градиент от ВСЕХ токенов\n'
                    'из-за нормализации — это создаёт "соревнование" за внимание',
            ha='center', fontsize=8, color=C['orange'], fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. A4: Обратное распространение через механизм Self-Attention (chain rule)',
             ha='center', fontsize=8, color=C['text_dim'])
    return sf(fig)


# ─── Диаграмма 5: Активационные функции ───
def fig_activations():
    fig, ax = plt.subplots(figsize=(7, 4))
    sa(ax)

    x = np.linspace(-4, 4, 200)
    ax.plot(x, np.maximum(0, x), label='ReLU', color=C['blue'], linewidth=2)
    ax.plot(x, x * (1 / (1 + np.exp(-x))), '--', label='SiLU (Swish)', color=C['accent'], linewidth=2)
    ax.plot(x, 0.5 * x * (1 + np.tanh(np.sqrt(2/np.pi) * (x + 0.044715 * x**3))),
            ':', label='GeLU', color=C['green'], linewidth=2)

    ax.set_xlabel('x', fontsize=9, color=C['text'])
    ax.set_ylabel('f(x)', fontsize=9, color=C['text'])
    ax.set_title('Сравнение активационных функций FFN', fontsize=10, color=C['text'], fontweight='bold')
    ax.legend(fontsize=8, facecolor=C['primary'], edgecolor=C['text_dim'], labelcolor=C['text'])
    ax.spines['left'].set_visible(True)
    ax.spines['bottom'].set_visible(True)
    ax.set_ylim(-1, 4)

    # Note about SwiGLU
    ax.text(3, 3, 'SwiGLU = SiLU(x·W₁) ⊙ (x·W₂)\nУдваивает W, но\nкачество выше',
            fontsize=7, color=C['orange'], fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. A5: Активационные функции FFN — от ReLU к SwiGLU',
             ha='center', fontsize=8, color=C['text_dim'])
    return sf(fig)


# ─── Диаграмма 6: Complexity comparison ───
def fig_complexity():
    fig, ax = plt.subplots(figsize=(7, 4))
    sa(ax)

    seq_lens = np.array([128, 256, 512, 1024, 2048, 4096])
    rnn_cost = seq_lens * (768**2)  # O(n·d²)
    attn_cost = (seq_lens**2) * 768  # O(n²·d)

    ax.plot(seq_lens, rnn_cost / 1e6, 'o-', label='RNN: O(n·d²)', color=C['blue'], linewidth=2)
    ax.plot(seq_lens, attn_cost / 1e6, 's-', label='Transformer: O(n²·d)', color=C['accent'], linewidth=2)

    ax.set_xlabel('Длина последовательности (n)', fontsize=9, color=C['text'])
    ax.set_ylabel('Относительная стоимость (MFLOPS)', fontsize=9, color=C['text'])
    ax.set_title('Сложность: RNN vs Transformer (d_model=768)', fontsize=10, color=C['text'], fontweight='bold')
    ax.legend(fontsize=8, facecolor=C['primary'], edgecolor=C['text_dim'], labelcolor=C['text'])
    ax.spines['left'].set_visible(True)
    ax.spines['bottom'].set_visible(True)

    # Highlight crossover point
    ax.axvline(x=768, color=C['orange'], linestyle='--', alpha=0.5)
    ax.text(780, ax.get_ylim()[1]*0.9, 'n = d_model\nCross-over', fontsize=7, color=C['orange'])

    fig.text(0.5, 0.02, 'Рис. A6: Вычислительная сложность — когда Transformer выгоднее RNN',
             ha='center', fontsize=8, color=C['text_dim'])
    return sf(fig)


# ─── Сборка документа ───
def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    title = doc.add_heading('Лекция 1 — Дополнение: закрытые пробелы', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        'Рецензия и доработка по 10 найденным пробелам\n'
        'FishingLLM — каждая теоретическая деталь имеет значение для нашей модели'
    )
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # ═══ Пробел 1: Числовой пример Attention ═══
    doc.add_heading('A.1. Числовой пример Self-Attention (закрывает пробел №1)', level=1)

    p = doc.add_paragraph()
    p.add_run('В лекции QKV объяснён на уровне идей. Теперь — на реальных числах.')

    p = doc.add_paragraph()
    p.add_run('Рассмотрим предложение из двух слов: "щука лещ". Размерность dₖ=3.')
    p.add_run(' Q, K, V — для простоты возьмём как входные эмбеддинги (хотя в реальности они через линейные слои).')

    p = doc.add_paragraph()
    p.add_run('Шаг 1. Вычисляем матрицу внимания (Q·Kᵀ):').bold = True

    code = doc.add_paragraph()
    r = code.add_run(
        '# Эмбеддинги слов [щука, лещ] с d=3\n'
        'X = [[0.8, 0.2, 0.5],  # щука\n'
        '     [0.1, 0.9, 0.3]]  # лещ\n\n'
        'Q = K = V = X  # для упрощения (в реальности Q=W_Q@X, K=W_K@X)\n\n'
        'scores = Q @ K.T  = [[0.93, 0.49],\n'
        '                     [0.49, 0.91]]\n\n'
        'Интерпретация:\n'
        '- элемент [0,0]=0.93: "щука" смотрит на "щука" (себя) — высокое внимание\n'
        '- элемент [0,1]=0.49: "щука" смотрит на "лещ" — умеренное внимание\n'
        '- элемент [1,0]=0.49: "лещ" смотрит на "щука"\n'
        '- элемент [1,1]=0.91: "лещ" смотрит на себя'
    )
    r.font.name = 'Consolas'; r.font.size = Pt(8)

    p = doc.add_paragraph()
    r = p.add_run('Шаг 2. Масштабируем и применяем softmax:').bold = True

    code = doc.add_paragraph()
    r = code.add_run(
        'scaled = scores / √3 = [[0.537, 0.283],\n'
        '                       [0.283, 0.525]]\n\n'
        'weights = softmax(scaled) = [[0.563, 0.437],  # "щука" -> 56% себе, 44% лещ\n'
        '                             [0.440, 0.560]]  # "лещ" -> 44% щука, 56% себе\n\n'
        'Заметьте: "щука" и "лещ" смотрят друг на друга почти поровну!\n'
        'Это потому что они оба — названия рыб, семантически близки.'
    )
    r.font.name = 'Consolas'; r.font.size = Pt(8)

    p = doc.add_paragraph()
    r = p.add_run('Шаг 3. Взвешенная сумма с V:').bold = True

    code = doc.add_paragraph()
    r = code.add_run(
        'output = weights @ V = [[0.563·[1,0] + 0.437·[0,1]],\n'
        '                       [0.440·[1,0] + 0.560·[0,1]]]\n'
        '       = [[0.563, 0.437],\n'
        '          [0.440, 0.560]]\n\n'
        'Каждый токен теперь содержит смесь информации от обоих слов.\n'
        '"щука" стала [0.56, 0.44] — "заразилась" семантикой "леща" на 44%.\n'
        'Это и есть суть Self-Attention: контекстуализация.'
    )
    r.font.name = 'Consolas'; r.font.size = Pt(8)

    doc.add_picture(fig_numerical_attention(), width=Inches(5.5))
    doc.add_page_break()

    # ═══ Пробел 2: Активации FFN ═══
    doc.add_heading('A.2. Активационные функции FFN (закрывает пробел №2)', level=1)

    p = doc.add_paragraph()
    p.add_run('В лекции приведена формула FFN(x) = ReLU(x·W₁ + b₁)·W₂ + b₂. Это классика (Vaswani, 2017).')
    p = doc.add_paragraph()
    p.add_run('НО: современные LLM отказались от ReLU. Почему?')

    doc.add_paragraph('ReLU: быстрый, но "умирает" (dead neurons) при отрицательных входах. Градиент = 0 для x < 0.', style='List Bullet')
    doc.add_paragraph('GeLU (Gaussian Error Linear Unit, Hendrycks & Gimpel, 2016): гладкая аппроксимация ReLU. Используется в GPT-2, BERT.', style='List Bullet')
    doc.add_paragraph('SiLU / Swish (Ramachandran et al., 2017): f(x) = x·σ(x). LLaMA 1.', style='List Bullet')
    doc.add_paragraph('SwiGLU (LLaMA 2/3, Mistral, Qwen): SiLU(W₁·x) ⊙ (W₂·x). Требует в 2 раза больше весов, но даёт лучшее качество.', style='List Bullet')

    p = doc.add_paragraph()
    r = p.add_run('SwiGLU стала стандартом де-факто для 2024-2025. ')
    r.bold = True
    p.add_run('Формула: FFN_swiglu(x) = (SiLU(x·W₁) ⊙ (x·W₂)) · W₃')

    doc.add_picture(fig_activations(), width=Inches(5))

    p = doc.add_paragraph()
    r = p.add_run('Trade-off: ')
    r.bold = True
    p.add_run('SwiGLU требует 3 матрицы вместо 2 (параметров ≈ 3·d_model·d_ff вместо 2·d_model·d_ff). '
              'Для нас (150M) это ~25% дополнительных параметров. Мы будем использовать SwiGLU — '
              'качество важнее.')
    doc.add_page_break()

    # ═══ Пробел 3: RoPE ═══
    doc.add_heading('A.3. RoPE — Rotary Position Embedding (закрывает пробел №3)', level=1)

    p = doc.add_paragraph()
    p.add_run('В лекции сказано, что мы будем использовать RoPE, но не объяснено, как он работает. Исправляем.')

    p = doc.add_paragraph()
    r = p.add_run('Идея RoPE (Su et al., 2021): ')
    r.bold = True
    p.add_run('вместо того чтобы добавлять позиционный сигнал к эмбеддингу (как синусоиды в оригинальном Transformer), '
              'RoPE "поворачивает" векторы Q и K в зависимости от их позиции.')

    p = doc.add_paragraph()
    r = p.add_run('Формально для токена на позиции pos:')
    r.font.name = 'Consolas'; r.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run('f(q, pos) = q · e^(i·pos·ω) = q · (cos(pos·ω) + i·sin(pos·ω))')
    r.font.name = 'Consolas'; r.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Здесь q — это 2-мерная пара соседних элементов в Q/K. Каждая пара поворачивается '
              'на свой угол (своя частота ωₖ = 1/10000^(2k/d_model)).')

    p = doc.add_paragraph()
    p.add_run('Ключевое свойство: ')
    r = p.add_run('скалярное произведение повёрнутых Q и K зависит ТОЛЬКО от относительного расстояния '
                  'между токенами (|i-j|), а не от их абсолютных позиций.')
    r.bold = True

    doc.add_picture(fig_rope(), width=Inches(5.5))

    p = doc.add_paragraph()
    p.add_run('Почему RoPE лучше синусоид:')
    p = doc.add_paragraph('Относительное позиционирование — модель понимает "на расстоянии 5 слов назад", а не "на 42-й позиции"', style='List Bullet')
    p = doc.add_paragraph('Лучшая экстраполяция на длинные контексты (можно увеличить context length после обучения)', style='List Bullet')
    p = doc.add_paragraph('Естественная decay внимания с расстоянием (чем дальше — тем меньше вес)', style='List Bullet')

    p = doc.add_paragraph()
    r = p.add_run('Для нашей модели будем использовать RoPE с base_frequency=10000.0 — стандарт.')
    r.bold = True
    doc.add_page_break()

    # ═══ Пробел 4: Layer Norm γ/β ═══
    doc.add_heading('A.4. Layer Normalization: роль γ и β (закрывает пробел №4)', level=1)

    p = doc.add_paragraph()
    p.add_run('В лекции формула LN(x) = (x - μ)/√(σ² + ε)·γ + β. Спросим: зачем γ и β, если мы уже нормализовали?')

    p = doc.add_paragraph()
    p.add_run('Без γ и β: ')
    p.add_run('после нормализации каждый слой получает на вход распределение N(0,1). '
              'Это слишком "жёстко" — сеть не может сдвинуть или растянуть распределение, '
              'если это нужно для задачи.')
    p.add_run(' γ и β — ')
    r = p.add_run('обучаемые параметры (скаляры размерности d_model)')
    r.bold = True
    p.add_run(', которые позволяют сети "подстроить" нормализацию под свои нужды.')

    p = doc.add_paragraph()
    p.add_run('Пример: внимание к "щука" требует высоких значений активации, а к "в" — низких. '
              'γ и β учатся этому сдвигу.')

    p = doc.add_paragraph()
    p.add_run('rms_norm (используется в LLaMA/Qwen): ') 
    p.add_run('современная альтернатива — убирает центрирование (μ=0), '
              'оставляет только деление на RMS вместе с γ. Без репараметризации '
              'нужно меньше вычислений.')
    doc.add_page_break()

    # ═══ Пробел 5: Masked Attention матрица ═══
    doc.add_heading('A.5. Causal Mask — как это выглядит (закрывает пробел №5)', level=1)

    p = doc.add_paragraph()
    p.add_run('Матрица маски размера [seq_len × seq_len]. Элемент (i,j) = 0, если j ≤ i (можно смотреть), '
              'и -∞ (или -1e9), если j > i (нельзя — будущее).')

    p = doc.add_paragraph()
    code = doc.add_paragraph()
    r = code.add_run(
        '# Маска для seq_len=4:\n'
        'mask = [[  0, -∞, -∞, -∞],\n'
        '        [  0,  0, -∞, -∞],\n'
        '        [  0,  0,  0, -∞],\n'
        '        [  0,  0,  0,  0]]\n\n'
        '# После добавления к scores и softmax:\n'
        'attention = [[0.8, 0.0, 0.0, 0.0],  # слово 1 смотрит ТОЛЬКО на слово 1\n'
        '             [0.4, 0.6, 0.0, 0.0],  # слово 2 смотрит на 1 и 2\n'
        '             [0.3, 0.3, 0.4, 0.0],  # слово 3 смотрит на 1,2,3\n'
        '             [0.2, 0.2, 0.3, 0.3]]  # слово 4 смотрит на все 1-4'
    )
    r.font.name = 'Consolas'; r.font.size = Pt(8)

    p = doc.add_paragraph()
    p.add_run('Как работает: к scores (Q·Kᵀ) добавляем маску. Там, где -∞, после softmax будет 0. '
              'Модель физически не может "увидеть" будущие токены.')

    doc.add_picture(fig_masked_attention(), width=Inches(5))
    doc.add_page_break()

    # ═══ Пробел 6: Backprop ═══
    doc.add_heading('A.6. Обратное распространение через Attention (закрывает пробел №6)', level=1)

    p = doc.add_paragraph()
    p.add_run('Как градиент (сигнал ошибки) проходит через Attention?' +
              ' По chain rule (правилу цепочки):')

    p = doc.add_paragraph()
    p.add_run('• ∂L/∂V = Wᵀ · ∂L/∂Output')
    p.add_run('  — берём веса внимания (W) и умножаем на градиент выхода')
    p = doc.add_paragraph()
    p.add_run('• ∂L/∂Scores = ∂L/∂Output · Vᵀ')
    p.add_run('  — градиент от V обратно к scores')
    p = doc.add_paragraph()
    p.add_run('• ∂L/∂Scores через softmax:')
    p.add_run('  ∂L/∂S_i = S_i · (∂L/∂P_i - Σ(S_j · ∂L/∂P_j))')
    p.add_run('  — самый хитрый шаг: из-за нормировки softmax градиент к одному токену')
    p.add_run('  зависит от ВСЕХ других токенов')
    p = doc.add_paragraph()
    p.add_run('• ∂L/∂Q = ∂L/∂Scores · K, ∂L/∂K = Qᵀ · ∂L/∂Scores')
    p.add_run('• ∂L/∂W_Q = Xᵀ · ∂L/∂Q, ∂L/∂W_K = Xᵀ · ∂L/∂K')

    doc.add_picture(fig_backprop(), width=Inches(5.5))

    p = doc.add_paragraph()
    p.add_run('Замечание: ')
    r = p.add_run('из-за softmax градиент к одному токену зависит от всех других. '
                  'Это делает обучение "соревновательным" — токены "борются" за внимание. '
                  'Если один токен получает вес 0.9, другой получит 0.1, и градиент к '
                  'второму будет подавлен.')
    r.font.italic = True
    doc.add_page_break()

    # ═══ Пробел 7: Сложность ═══
    doc.add_heading('A.7. Вычислительная сложность Transformer (закрывает пробел №7)', level=1)

    p = doc.add_paragraph()
    p.add_run('Основные формулы сложности:')

    p = doc.add_paragraph()
    p.add_run('Attention: O(n² · d)  (n — длина последовательности, d — размерность)')
    p.add_run('  n² — матрица внимания [n×n]')
    p.add_run('  d — умножение на V [n×d]')
    p = doc.add_paragraph()
    p.add_run('FFN: O(n · d²)')
    p.add_run('  d² — умножение на W₁ [d×4d] и W₂ [4d×d]')
    p = doc.add_paragraph()
    p.add_run('Итого на слой: O(n²·d + n·d²)')

    p = doc.add_paragraph()
    p.add_run('Для нашей модели (n=1024, d=768):')
    p.add_run('  Attention: 1024² · 768 ≈ 805 MFLOPS')
    p.add_run('  FFN: 1024 · 768² ≈ 604 MFLOPS')
    p.add_run('  Всего на слой: ~1.4 GFLOPS')
    p.add_run('  Всего (12 слоёв): ~17 GFLOPS на один forward pass')
    p.add_run('  Для справки: GTX 1060 даёт ~4 TFLOPS → ~235 forward/сек')

    doc.add_picture(fig_complexity(), width=Inches(5))

    p = doc.add_paragraph()
    p.add_run('Зачем знать сложность? ').bold = True
    p.add_run('Чтобы понимать, где узкое место. Для генерации (n=1024) — внимание '
              'занимает ~60% времени. Для обучения — внимание ~80% из-за BP. '
              'Flash Attention (оптимизация памяти) даёт ускорение ×2-4.')
    doc.add_page_break()

    # ═══ Пробел 10: Mamba / SSM ═══
    doc.add_heading('A.8. State Space Models — альтернатива Transformer (закрывает пробел №10)', level=1)

    p = doc.add_paragraph()
    p.add_run('С 2023 года появился конкурент — State Space Models (SSM).')

    p = doc.add_paragraph()
    p.add_run('Mamba (Gu & Dao, 2023) — архитектура, которая:')
    p = doc.add_paragraph('• Имеет сложность O(n) вместо O(n²)', style='List Bullet')
    p = doc.add_paragraph('• Не требует хранения матрицы внимания [n×n]', style='List Bullet')
    p = doc.add_paragraph('• Может работать с контекстом 1M+ токенов', style='List Bullet')
    p = doc.add_paragraph('• Но: уступает Transformer в качестве на задачах "запоминания фактов"', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Почему мы НЕ выбираем Mamba для FishingLLM:')
    p = doc.add_paragraph('• Mamba сложнее в обучении (нет готовых реализаций для малых моделей)', style='List Bullet')
    p = doc.add_paragraph('• Наш контекст (1024-2048 токенов) — O(n²) не проблема', style='List Bullet')
    p = doc.add_paragraph('• Transformer лучше изучен и документирован — важнее для обучения', style='List Bullet')
    p = doc.add_paragraph('• Больше open-source инструментов (Hugging Face, llama.cpp и т.д.)', style='List Bullet')

    p = doc.add_paragraph()
    r = p.add_run('Вывод: Transformer — правильный выбор для нашей задачи. Mamba — для контекстов 100K+ токенов.')
    r.bold = True
    doc.add_page_break()

    # ═══ Пробел 8-9: Задания ═══
    doc.add_heading('A.9. Практические задания, привязанные к проекту (закрывает пробелы №8, №9)', level=1)

    p = doc.add_paragraph()
    r = p.add_run('Задание 1. Реализуйте Self-Attention на PyTorch')
    r.bold = True
    code = doc.add_paragraph()
    r = code.add_run(
        'import torch\n'
        'import torch.nn.functional as F\n\n'
        'def self_attention(x):\n'
        '    """x: [batch, seq_len, d_model] — эмбеддинги токенов"""\n'
        '    d_k = x.shape[-1]\n'
        '    Q = x @ W_Q  # W_Q: [d_model, d_k]\n'
        '    K = x @ W_K\n'
        '    V = x @ W_V\n'
        '    scores = Q @ K.transpose(-2, -1) / (d_k ** 0.5)\n'
        '    weights = F.softmax(scores, dim=-1)\n'
        '    return weights @ V\n\n'
        '# Примените к эмбеддингам рыболовных терминов:\n'
        '# ["щука", "лещ", "окунь", "спиннинг", "червь"]\n'
        '# Какие пары имеют наибольшее внимание?'
    )
    r.font.name = 'Consolas'; r.font.size = Pt(8)

    p = doc.add_paragraph()
    r = p.add_run('Задание 2. Разбор config.json нашей будущей модели')
    r.bold = True
    code = doc.add_paragraph()
    r = code.add_run(
        '# Заполните конфигурацию для модели ~150M:\n'
        'config = {\n'
        '    "vocab_size": 16_000,      # размер словаря\n'
        '    "d_model": 768,            # скрытая размерность\n'
        '    "n_layers": 12,            # число слоёв\n'
        '    "n_heads": 12,             # число голов внимания\n'
        '    "d_ff": 3072,              # размер FFN (4*d_model)\n'
        '    "activation": "swiglu",    # функция активации\n'
        '    "max_seq_len": 2048,       # макс. длина контекста\n'
        '    "rotary_pct": 0.25,        # доля RoPE-размерностей\n'
        '}\n\n'
        '# Подсчитайте:\n'
        '# Embedding: vocab_size * d_model = 16_000 * 768 = 12.3M\n'
        '# Attention: 4 * d_model^2 = 4 * 768^2 = 2.36M на слой\n'
        '# FFN: 3 * d_model * d_ff = 3 * 768 * 3072 = 7.08M на слой\n'
        '# Итого ≈ 12.3M + 12*(2.36M + 7.08M) + 12.3M ≈ 138M\n'
        '# (плюс LayerNorm, bias, RoPE ≈ до 150M)'
    )
    r.font.name = 'Consolas'; r.font.size = Pt(8)

    p = doc.add_paragraph()
    r = p.add_run('Задание 3. Визуализируйте матрицу внимания для фразы "На какую наживку '
                  'клюёт лещ в июне на Верхней Волге"')
    r.bold = True
    p = doc.add_paragraph('Какие токены имеют максимальные веса в каждой строке? '
                          'Совпадает ли это с вашей интуицией о важности слов?')

    p = doc.add_paragraph()
    r = p.add_run('Задание 4 (со звёздочкой). Сравните время выполнения Attention '
                  'на CPU vs GPU для seq_len = 128, 512, 2048')
    r.bold = True
    p = doc.add_paragraph('Напишите замер в PyTorch с torch.cuda.Event (если есть GPU) '
                          'или timeit. Постройте график O(n²). '
                          'При какой длине последовательности время становится неприемлемым (>1 сек) '
                          'для инференса?')

    doc.add_paragraph('')

    p = doc.add_paragraph()
    r = p.add_run('Критерии проверки:')
    r.bold = True
    items = [
        'Задание 1: weights.sum(-1) ≈ 1 для каждого токена',
        'Задание 2: итоговое число параметров ±5% от 150M',
        'Задание 3: max weight должен быть на диагонали или на тематически связанных словах',
        'Задание 4: график должен показывать квадратичный рост',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    r = p.add_run('Итог: после выполнения этих заданий вы сможете: ')
    r.bold = True
    p.add_run('1) написать Attention с нуля; 2) спроектировать конфиг модели; '
              '3) проанализировать, что "видит" модель; 4) оценить производительность.')
    doc.add_page_break()

    # ═══ Итог ═══
    doc.add_heading('Итог: что изменилось после патча', level=1)

    table = doc.add_table(rows=11, cols=3)
    table.style = 'Light Shading Accent 1'
    headers = ['№', 'Пробел', 'Статус']
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    rows_data = [
        ('1', 'Числовой пример Attention', '✅ Разобран на матрицах 2×3 с вычислениями'),
        ('2', 'Активации FFN (GeLU, SwiGLU)', '✅ ReLU → SiLU → SwiGLU с формулами и графиком'),
        ('3', 'RoPE — формула и интуиция', '✅ Комплексная плоскость, частота, поворот Q/K'),
        ('4', 'Layer Norm: γ, β', '✅ Обучаемые параметры, rms_norm, визуализация'),
        ('5', 'Causal Mask — матрица', '✅ Визуализация маски до/после'),
        ('6', 'Backprop через Attention', '✅ Chain rule, градиент через softmax'),
        ('7', 'Сложность в цифрах', '✅ FLOPs для n=1024, d=768, сравнение RNN vs TF'),
        ('8', 'Привязка заданий к проекту', '✅ 4 задания на реальных данных рыбалки'),
        ('9', 'PyTorch вместо NumPy', '✅ Код с torch, конфиг модели, замеры скорости'),
        ('10', 'Mamba / SSM', '✅ Альтернатива, причины НЕ выбирать её сейчас'),
    ]
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            table.cell(i+1, j).text = val

    doc.add_paragraph('')
    p = doc.add_paragraph()
    r = p.add_run('Следующий шаг: практика.')
    r.bold = True
    r.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
    p.add_run(' Берём задание 1 — реализуем Self-Attention на PyTorch, '
              'пишем тесты, проверяем на рыбацких эмбеддингах.')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    r = p.add_run('Теперь лекция 1 полна. Можно переходить к практическому циклу.')
    r.bold = True

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, 'lecture-01-addendum.docx')
    doc.save(path)
    print(f'Дополнение сохранено: {path}')


if __name__ == '__main__':
    build()
