from docx import Document
from docx.shared import Pt, Cm, RGBColor
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.Name = "Times New Roman"
font.Size = Pt(12)

section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    return h

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def add_code_block(code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

# =============== TITLE ===============
title = doc.add_heading("Практическая реализация механизма внимания", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

add_para("Документация кода: Self-Attention, Multi-Head Attention и GQA для проекта FishingLLM")
add_para("")

# =============== 1. Overview ===============
add_heading("1. Общая схема работы attention")

add_para(
    "Прежде чем разбирать код, посмотрим на поток данных. "
    "Входной тензор x размерности (batch, seq_len, d_model) проходит "
    "через три независимые проекции, затем через attention, "
    "сборку голов и обратную проекцию:"
)

add_code_block(
    "x  (B, L, d_model)\n"
    "|\n"
    "+-- W_q --> Q (B, n_heads, L, d_k) --+\n"
    "|                                      |\n"
    "+-- W_k --> K (B, n_kv, L, d_k) ----> Attention(Q,K,V)\n"
    "|                                      |\n"
    "+-- W_v --> V (B, n_kv, L, d_k) ---->  |\n"
    "                                       |\n"
    "                               concat heads\n"
    "                                       |\n"
    "                                      W_o\n"
    "                                       |\n"
    "                               output (B, L, d_model)"
)

add_para(
    "Зачем нужны три матрицы W_q, W_k, W_v? Если бы мы вычисляли внимание "
    'напрямую от x к x, каждый токен "узнавал" бы только самого себя. '
    "Три проекции позволяют модели разделить роли: Q решает, что искать, "
    "K — что предлагать, V — что передавать. Это как библиотека: "
    "вы приходите с запросом (Q), библиотекарь сверяет каталог (K) "
    "и выдаёт книгу (V). Без этого разделения запрос и каталог были бы "
    "одним и тем же."
)

# =============== 2 ===============
add_heading("2. ScaledDotProductAttention — одно головое внимание")

add_para(
    "Класс ScaledDotProductAttention реализует формулу из Лекции 1: "
    "Attention(Q,K,V)=softmax(Q*K^T/sqrt(d_k))*V. "
    "Разберём forward построчно."
)

add_code_block("def forward(self, Q, K, V, mask=None):\n"
    "    scores = torch.matmul(Q, K.transpose(-2, -1)) * self.scale")

add_para(
    "torch.matmul(Q, K.transpose(-2, -1)) вычисляет Q*K^T. "
    "transpose(-2, -1) меняет две последние оси K: "
    "из (batch, seq_len, d_k) получаем (batch, d_k, seq_len). "
    "Умножение даёт (batch, seq_len, seq_len). "
    "self.scale = 1/sqrt(d_k) — нормировка, предотвращающая "
    "взрыв значений softmax при больших d_k."
)

add_code_block("if mask is not None:\n    scores = scores.masked_fill(mask, float('-inf'))")

add_para(
    "Маска — тензор bool формы (batch, seq_len, seq_len). "
    "Где mask=True, ставится -inf. После softmax при наличии "
    "хотя бы одного незамаскированного элемента замаскированные "
    "позиции получают вероятность ровно 0. "
    "Для каузальной маски запрещаем токену attendить к будущим."
)

add_code_block("attn = F.softmax(scores, dim=-1)\nreturn torch.matmul(attn, V), attn")

add_para(
    "Softmax по последней оси превращает scores в распределение "
    "вероятностей. Умножение attn на V даёт взвешенную сумму "
    "значений. Веса attn возвращаются для отладки."
)

# =============== 3. GQA Motivation ===============
add_heading("3. Проблема памяти: почему MHA дорог")

add_para(
    "В классическом Multi-Head Attention (MHA) каждая голова имеет "
    "собственные K и V. При авторегрессивной генерации все K и V "
    "предыдущих токенов хранятся в KV-кэше. Для 12 голов, d_k=64, "
    "контекст 2048: 2*2048*12*64*2 байта = 6 МБ на слой, 72 МБ "
    "на все 12 слоёв. На CPU с 8GB RAM это терпимо, но при "
    "контексте 4096+ становится проблемой."
)

add_para(
    "Grouped Query Attention (GQA, Ainslie et al., 2023) — компромисс: "
    "n_kv_heads < n_heads. Группа Q-голов делит одни K и V."
)

add_para("Сравнение подходов:", bold=True)

add_para(
    "- MQA (Shazeer, 2019): n_kv_heads = 1. Максимальная экономия "
    "кэша, но падение качества.\n"
    "- MHA (оригинал): n_kv_heads = n_heads. Максимум качества, "
    "но дорогой inference.\n"
    "- GQA: промежуточный вариант. Для 12:4 — трёхкратное сжатие "
    "KV-кэша при <1% потери качества."
)

add_para(
    "Для FishingLLM на CPU каждый сэкономленный мегабайт позволяет "
    "обрабатывать более длинные промпты. При 4 KV-головах кэш "
    "составляет 24 МБ на 12 слоёв против 72 МБ в полном MHA."
)

# =============== 4 ===============
add_heading("4. MultiHeadAttention — конструктор")

add_para(
    "Параметры: d_model — размерность эмбеддингов (768); "
    "n_heads — число голов (12); n_kv_heads — KV-голов (4); "
    "dropout — вероятность обнуления. "
    "assert проверяет, что d_model делится на n_heads."
)

add_para("Четыре линейные проекции:", bold=True)
add_para(
    "- W_q: d_model -> n_heads * d_k — запросы (все головы)\n"
    "- W_k: d_model -> n_kv_heads * d_k — ключи (только KV-головы)\n"
    "- W_v: d_model -> n_kv_heads * d_k — значения\n"
    "- W_o: n_heads * d_k -> d_model — обратная проекция"
)

add_para(
    "bias=False — как в LLaMA и GPT. Сдвиги вносятся нормализацией "
    "(RMSNorm), поэтому bias в проекциях избыточен."
)

add_para(
    "n_rep = n_heads // n_kv_heads — сколько Q-голов приходится "
    "на одну KV-голову. При 12:4, n_rep=3."
)

add_heading("4.1 Инициализация и маска", level=3)

add_para(
    "Xavier Uniform — распространённый выбор для линейных слоёв "
    "в attention. Он сохраняет дисперсию сигнала при проходе."
)

add_code_block("@staticmethod\n"
    "def _create_causal_mask(seq_len, device):\n"
    "    mask = torch.triu(torch.full((seq_len, seq_len), True,\n"
    "        dtype=torch.bool, device=device), diagonal=1)\n"
    "    return mask.unsqueeze(0)")

add_para(
    "torch.triu с diagonal=1 даёт True для элементов выше "
    "главной диагонали: токен i attendит только к j <= i. "
    "Маска кэшируется — при повторном вызове с той же длиной "
    "не создаётся заново. unsqueeze(0) добавляет batch-размерность."
)

# =============== 5 ===============
add_heading("5. Forward-метод MultiHeadAttention")

add_para("Шаг 1. Проекции Q, K, V:", bold=True)

add_code_block("Q = self.W_q(x).view(B, L, self.n_heads, self.d_k).transpose(1, 2)\n"
    "K = self.W_k(x).view(B, L, self.n_kv_heads, self.d_k).transpose(1, 2)\n"
    "V = self.W_v(x).view(B, L, self.n_kv_heads, self.d_k).transpose(1, 2)")

add_para(
    "После view Q — (B, L, n_heads, d_k), после transpose(1,2) — "
    "(B, n_heads, L, d_k). K и V — (B, n_kv_heads, L, d_k). "
    "Переставляем оси, чтобы размерность L была предпоследней."
)

add_para("Шаг 2. Повтор KV-голов для GQA:", bold=True)

add_code_block("if self.n_rep > 1:\n"
    "    K = K.repeat_interleave(self.n_rep, dim=1)\n"
    "    V = V.repeat_interleave(self.n_rep, dim=1)")

add_para(
    "repeat_interleave копирует каждую KV-голову n_rep раз. "
    "Форма: (B, n_kv_heads, L, d_k) -> (B, n_heads, L, d_k). "
    "Теперь K и V согласованы с Q по числу голов."
)

add_para("Шаг 3. Attention ко всем головам:", bold=True)

add_code_block("attn_output, attn_weights = self.attn(\n"
    "    Q.reshape(-1, L, self.d_k),\n"
    "    K.reshape(-1, L, self.d_k),\n"
    "    V.reshape(-1, L, self.d_k), mask)")

add_para(
    "reshape(-1, L, d_k) объединяет batch и n_heads: "
    "(B * n_heads, L, d_k). Это превращает multi-head "
    "в batch-обработку независимых single-head attention."
)

add_para("Шаг 4. Обратная проекция:", bold=True)

add_code_block(
    "attn_output = attn_output.view(B, self.n_heads, L, self.d_k)\n"
    "    .transpose(1, 2).reshape(B, L, -1)\n"
    "attn_output = self.dropout(self.W_o(attn_output))")

add_para(
    "view восстанавливает головы, transpose возвращает "
    "(B, L, n_heads, d_k), reshape собирает головы в последнюю "
    "ось. W_o проецирует n_heads * d_k обратно в d_model."
)

# =============== 6 ===============
add_heading("6. Почему attention квадратичный?")

add_para(
    "Матрица scores имеет размер (L, L): каждый токен "
    "вычисляет скалярное произведение со всеми L токенами. "
    "Сложность O(L^2 * d_k) — квадратичная по длине контекста. "
    "Это фундаментальное ограничение vanilla attention: "
    "при L=2048 это ~4M операций на голову, при L=8192 — ~67M. "
    "Отсюда мотивация для FlashAttention, Linear Attention, "
    "Sliding Window — они снижают сложность до O(L) или O(L*log(L))."
)

# =============== 7 ===============
add_heading("7. Тесты")

add_para(
    "Тесты (tests/test_attention.py) разделены на два класса: "
    "TestScaledDotProductAttention (7 тестов) и "
    "TestMultiHeadAttention (11 тестов). "
    "Все 18 тестов проходят."
)

add_heading("7.1 ScaledDotProductAttention", level=3)

add_para(
    "test_output_shape — проверка форм: выход (B, L, d_k), "
    "веса (B, L, L).")
add_para(
    "test_attention_weights_sum_to_one — свойство softmax: "
    "сумма весов для каждого токена = 1.")
add_para(
    "test_causal_mask_prevents_future_attention — weights[i,j]=0 "
    "для j>i. Гарантия авторегрессионности.")
add_para(
    "test_identical_qkv_no_nan — при Q=K=V нет NaN, "
    "формы корректны (замена хрупкого теста diag > 0.5).")
add_para(
    "test_fully_masked_row_no_nan — строка, где все позиции "
    "замаскированы, не даёт NaN.")
add_para(
    "test_scaling_factor — масштабирование sqrt(d_k) "
    "предотвращает NaN при больших значениях.")
add_para(
    "test_gradients_flow — градиенты проходят через Q, K, V, "
    "не содержат NaN или None.")

add_heading("7.2 MultiHeadAttention", level=3)

add_para(
    "test_output_shape — (B, L, d_model) на выходе.")
add_para(
    "test_gqa_kv_repeat — n_rep=2 для 8:4, форма attention "
    "(B, n_heads, L, L).")
add_para(
    "test_causal_mask_works — нет NaN с маской.")
add_para(
    "test_without_gqa — n_kv_heads = n_heads, если не задан.")
add_para(
    "test_invalid_gqa_config_raises — AssertionError при "
    "n_heads % n_kv_heads != 0 (например, 10:3).")
add_para(
    "test_projection_changes_representation — выход не нулевой "
    "и не равен входу.")
add_para(
    "test_different_lengths — работает для 5 и 20 токенов.")
add_para(
    "test_return_attn_false — возвращает только тензор, "
    "не кортеж.")
add_para(
    "test_user_mask_batch — пользовательская маска (B, L, L) "
    "с B > 1 правильно транслируется на головы.")
add_para(
    "test_gqa_vs_vanilla_equivalence_with_shared_weights — "
    "GQA и full MHA дают численно совпадающий выход при "
    "скопированных в правильном порядке весах.")
add_para(
    "test_gradients_flow — все параметры имеют градиенты "
    "без NaN.")

# =============== 8 ===============
add_heading("8. Trade-offs и альтернативы")

add_para(
    "1. GQA (12:4) vs MHA (12:12): выбрана GQA ради "
    "экономии KV-кэша на CPU. MQA (12:1) — ещё больше "
    "экономия, но грубее.")
add_para(
    '2. bias=False: ~300K параметров меньше для d_model=768, '
    '12 голов. Стандарт LLaMA.')
add_para(
    "3. Xavier — распространённый выбор для attention. "
    "Kaiming предпочтительнее для ReLU-слоёв.")
add_para(
    "4. Маска генерируется на лету с кэшированием: "
    "универсально для разной длины.")
add_para(
    "5. Все головы в один batch через reshape(-1, L, d_k): "
    "векторизация, до 10x быстрее цикла по головам на GPU.")

# =============== 9 ===============
add_heading("9. Связь с теорией")

add_para(
    "Каждая строка кода соотносится с формулами из Лекции 1:"
)

add_para(
    "- Attention(Q,K,V)=softmax(Q*K^T/sqrt(d_k))*V — "
    "ScaledDotProductAttention.forward.\n"
    "- Multi-head: параллельные головы с формулой "
    "MultiHead = Concat(head_1,...,head_h)*W_o.\n"
    "- Каузальная маска — авторегрессионность decoder-only.\n"
    "- W_o — Concat(...)*W_o в лекционной нотации."
)

add_para(
    "RoPE-позиционное кодирование, residual connections "
    "и LayerNorm будут добавлены в следующих циклах "
    "на уровне блока трансформера."
)

# =============== 10 ===============
add_heading("10. Производительность (грубая оценка)")

add_para(
    "Для модели 150M, d_model=768, 12 голов, d_k=64, "
    "контекст 2048 (порядок величины, зависит от CPU, BLAS, "
    "количества потоков):"
)

add_para(
    "- Q*K^T: O(B*L^2*d_k) ~ 2*2048^2*64 ~ 537M оп/слой.\n"
    "- attn*V: те же 537M.\n"
    "- Итого на 12 слоёв: ~12.9 млрд оп/forward.\n"
    "- Скорость на CPU: ~1-2 сек на токен (грубая оценка).\n"
    "- KV-кэш: 2*2048*4*64*2 = 2 МБ/слой, 24 МБ всего.\n"
    "- Без GQA (12 голов): 72 МБ — разница в 3 раза."
)

add_para(
    "Для контекста 4096+ разница становится ещё заметнее: "
    "GQA даёт ~96 МБ против ~288 МБ в полном MHA."
)

out_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "lecture-attention-implementation.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Size: {os.path.getsize(out_path)} bytes")
