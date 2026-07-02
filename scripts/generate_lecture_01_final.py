"""Лекция 1 (полная): Transformer — единая версия, включающая все патчи
Источники: Vaswani 2017, Su 2021, Touvron 2023, Shazeer 2020, Kaplan 2020
"""

import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Arc
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io, os, sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
OUTPUT = os.path.join(os.path.dirname(__file__), '..', 'docs')

C = {'bg':'#1a1a2e','primary':'#16213e','secondary':'#0f3460','accent':'#e94560',
     'orange':'#f59e0b','green':'#10b981','blue':'#3b82f6','purple':'#8b5cf6',
     'teal':'#06b6d4','pink':'#ec4899','text':'#e2e8f0','text_dim':'#94a3b8'}

def sa(ax):
    ax.set_facecolor(C['bg']); ax.figure.patch.set_facecolor(C['bg'])
    for s in ax.spines.values(): s.set_visible(False)
    ax.tick_params(colors=C['text_dim'],labelsize=6)
    ax.grid(True,alpha=0.06,color=C['text_dim'])
def sf(fig):
    buf=io.BytesIO(); fig.savefig(buf,format='png',dpi=160,bbox_inches='tight',facecolor=C['bg']); buf.seek(0); plt.close(fig); return buf

# ─── FIG 1: RNN → LSTM → Transformer ───
def fig_history():
    fig,axes=plt.subplots(1,3,figsize=(11,3.5))
    titles=['RNN (1986)','LSTM (1997)','Transformer (2017)']
    descs=[['hₜ = tanh(W·[hₜ₋₁;xₜ])','O(n) sequential','vanishing gradient'],
           ['Cₜ = fₜ·Cₜ₋₁ + iₜ·C̃ₜ','resolves vanishing grad','still sequential O(n)'],
           ['Attention(Q,K,V)','all tokens at once O(1)','scales to 100B+ params']]
    for idx,(ax,title,desc) in enumerate(zip(axes,titles,descs)):
        sa(ax); ax.axis('off')
        ax.text(0.5,0.85,title,ha='center',fontsize=10,color='white',fontweight='bold',transform=ax.transAxes)
        for j,d in enumerate(desc):
            c=[C['text'],C['text'],C['accent']][j]
            ax.text(0.5,0.55-j*0.2,d,ha='center',fontsize=7.5,color=c,transform=ax.transAxes)
        if idx<2:# draw arrow
            ax.annotate('',xy=(0.86,0.5),xytext=(0.14,0.5),arrowprops=dict(arrowstyle='->',color=C['accent'],lw=2),transform=ax.transAxes)
    fig.text(0.5,0.02,'Рис.1: Эволюция архитектур для последовательностей — от RNN к Transformer',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 2: Attention numerical example ───
def fig_attention_numerical():
    fig,axes=plt.subplots(1,3,figsize=(11,3.5))
    Q=np.array([[0.8,0.2,0.5],[0.1,0.9,0.3]]); K=Q.copy(); V=Q.copy()
    scores=Q@K.T; scaled=scores/np.sqrt(3)
    exp_s=np.exp(scaled-scaled.max(axis=1,keepdims=True)); weights=exp_s/exp_s.sum(axis=1,keepdims=True)
    output=weights@V
    mats=[scores,weights,output]; titles=['Шаг 1: Q·Kᵀ','Шаг 2: softmax(Q·Kᵀ/√3)','Шаг 3: взвешенная сумма']
    for idx,(ax,mat,title) in enumerate(zip(axes,mats,titles)):
        sa(ax); ax.axis('off')
        ax.text(0.5,0.92,title,ha='center',fontsize=9,color='white',fontweight='bold',transform=ax.transAxes)
        if idx==0: labels=['щука','лещ']
        else: labels=['','']
        n_cols=2 if mat.shape[1]==2 else 3
        for i in range(2):
            for j in range(n_cols):
                v=mat[i,j]; mx=mat.max(); mn=mat.min(); rng_v=mx-mn+1e-8
                c_use=C['accent'] if v==mx else C['blue'] if v>mat.mean() else C['teal']
                a=0.3+0.6*(v-mn)/rng_v
                rect=FancyBboxPatch((j*0.22+0.08,1.0-(i+0.65)/2),0.18,0.18,boxstyle="round,pad=0.01",facecolor=c_use,alpha=a,edgecolor='white',linewidth=0.3)
                ax.add_patch(rect)
                ax.text(j*0.22+0.17,1.0-(i+0.56)/2,f'{v:.3f}',ha='center',va='center',fontsize=7,color='white')
            if idx==0: ax.text(0.01,1.0-(i+0.5)/2,labels[i],fontsize=8,color=C['text'])
    fig.text(0.5,0.02,'Рис.2: Self-Attention для "щука лещ" (d=3) — от сырых scores до контекстуализированного выхода',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 3: QKV formula breakdown ───
def fig_qkv():
    fig,ax=plt.subplots(figsize=(9,5)); sa(ax); ax.axis('off')
    ax.text(2.5,2.9,'Attention(Q,K,V) = softmax( Q·Kᵀ / √dₖ ) · V',ha='center',fontsize=13,color=C['orange'],fontweight='bold',family='monospace')
    boxes=[('Q: [n×dₖ]','"что я ищу?"','каждый токен формирует запрос к остальным',C['blue']),
           ('K: [n×dₖ]','"что у меня есть?"','каждый токен отвечает, насколько он релевантен',C['purple']),
           ('Q·Kᵀ: [n×n]','матрица внимания','скалярное произведение → сила связи между парами',C['teal']),
           ('÷√dₖ','масштабирование','не даёт softmax стать too sharp при большой dₖ',C['green']),
           ('softmax','вероятности','сумма по строке = 1, соревнование за внимание',C['accent']),
           ('·V: [n×dᵥ]','взвешенная сумма','каждый токен = смесь информации от всех',C['pink'])]
    for i,(title,q,desc,c) in enumerate(boxes):
        x=0.2+(i%3)*2.8; y=1.8-(i//3)*1.0
        r=FancyBboxPatch((x,y),2.4,0.55,boxstyle="round,pad=0.03",facecolor=c,alpha=0.15,edgecolor=c,linewidth=1.5); ax.add_patch(r)
        ax.text(x+1.2,y+0.4,title,ha='center',va='center',fontsize=8,color='white',fontweight='bold')
        ax.text(x+1.2,y+0.2,q,ha='center',va='center',fontsize=7,color=C['orange'])
        ax.text(x+1.2,y+0.05,desc,ha='center',va='top',fontsize=6.5,color=C['text'])
    fig.text(0.5,0.02,'Рис.3: Разбор формулы Attention — символ за символом',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 4: RoPE ───
def fig_rope_full():
    fig,ax=plt.subplots(figsize=(8,4.5)); sa(ax); ax.axis('off')
    # complex plane with rotations
    ax.text(0.5,-0.05,'RoPE: f_q(qₘ,m) = qₘ · e^(i·m·θ)\nгде θₖ = 10000^(-2k/d)',ha='center',fontsize=9,color=C['orange'],fontweight='bold',family='monospace',transform=ax.transAxes)
    # left: pos 0
    ax1=plt.axes([0.05,0.28,0.38,0.6]); sa(ax1)
    ax1.set_xlim(-1.5,1.5); ax1.set_ylim(-1.5,1.5); ax1.set_aspect('equal')
    circ=plt.Circle((0,0),1.0,fill=False,color=C['text_dim'],alpha=0.2,linewidth=0.5); ax1.add_patch(circ)
    ax1.arrow(0,0,1,0,head_width=0.1,head_length=0.1,fc=C['blue'],ec=C['blue'])
    ax1.text(0.5,0.15,'Q₀',fontsize=9,color=C['blue'],fontweight='bold')
    ax1.set_title('Позиция 0: Q₀ = q',fontsize=8,color=C['text'])
    ax1.text(0,-1.35,'rotated by 0 rad',fontsize=6,color=C['text_dim'],ha='center')
    # right: rotated
    ax2=plt.axes([0.55,0.28,0.38,0.6]); sa(ax2)
    ax2.set_xlim(-1.5,1.5); ax2.set_ylim(-1.5,1.5); ax2.set_aspect('equal')
    circ2=plt.Circle((0,0),1.0,fill=False,color=C['text_dim'],alpha=0.2,linewidth=0.5); ax2.add_patch(circ2)
    angle=np.pi/3
    vec=np.array([np.cos(angle),np.sin(angle)])
    ax2.arrow(0,0,1,0,head_width=0.05,head_length=0.05,fc=C['text_dim'],ec=C['text_dim'],alpha=0.3)
    ax2.arrow(0,0,vec[0],vec[1],head_width=0.1,head_length=0.1,fc=C['accent'],ec=C['accent'])
    ax2.text(vec[0]*0.5,vec[1]*0.5+0.1,'Qₘ',fontsize=9,color=C['accent'],fontweight='bold')
    arc=Arc((0,0),0.6,0.6,angle=0,theta1=0,theta2=np.degrees(angle),color=C['orange'],linewidth=1); ax2.add_patch(arc)
    ax2.text(0.2,0.1,f'm·θ',fontsize=7,color=C['orange'])
    ax2.set_title(f'Позиция m: Qₘ = q·e^(i·m·θ)',fontsize=8,color=C['text'])
    ax2.text(0,-1.35,f'rotated by {angle:.1f} rad',fontsize=6,color=C['text_dim'],ha='center')
    # property
    ax.text(0.5,0.88,'Свойство RoPE: ⟨Qₘ, Kₙ⟩ = ⟨Q₀, Kₙ₋ₘ⟩\nСкалярное произведение зависит ТОЛЬКО от РАЗНОСТИ позиций',ha='center',fontsize=9,color=C['green'],fontweight='bold',transform=ax.transAxes)
    fig.text(0.5,0.02,'Рис.4: Rotary Position Embedding (Su et al., 2021)',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 5: Decoder Layer ───
def fig_decoder():
    fig,ax=plt.subplots(figsize=(5,7)); sa(ax); ax.axis('off')
    layers=[(-0.35,1.1,0.7,0.1,'ВХОД (X+PE)',C['green'],True),
            (-0.30,0.85,0.6,0.1,'RMS Norm',C['teal'],True),
            (-0.35,0.55,0.7,0.18,'GQA\nGrouped Query Attn',C['purple'],True),
            (-0.30,0.25,0.6,0.1,'RMS Norm',C['teal'],True),
            (-0.35,-0.05,0.7,0.18,'SwiGLU FFN',C['blue'],True),
            (-0.25,-0.35,0.5,0.1,'ВЫХОД (→ слой N+1)',C['accent'],True)]
    for x,y,w,h,lbl,clr,_ in layers:
        r=FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.03",facecolor=clr,alpha=0.85,edgecolor='white',linewidth=0.5); ax.add_patch(r)
        ax.text(x+w/2,y+h/2,lbl,ha='center',va='center',fontsize=7,color='white',fontweight='bold')
    # arrows
    for i in range(len(layers)-1):
        ax.annotate('',xy=(0,layers[i+1][1]),xytext=(0,layers[i][1]),
                    arrowprops=dict(arrowstyle='->',color=C['text'],lw=1.5))
    # residuals
    ax.annotate('',xy=(0.38,0.55),xytext=(0.38,1.1),arrowprops=dict(arrowstyle='->',color=C['orange'],lw=1.5,linestyle='dashed'))
    ax.text(0.42,0.82,'+',fontsize=10,color=C['orange'],fontweight='bold')
    ax.annotate('',xy=(0.38,-0.05),xytext=(0.38,0.55),arrowprops=dict(arrowstyle='->',color=C['orange'],lw=1.5,linestyle='dashed'))
    ax.text(0.42,0.25,'+',fontsize=10,color=C['orange'],fontweight='bold')
    ax.text(0.65,0.85,'x_out = x_in + Layer(x_in)',fontsize=6,color=C['orange'])
    fig.text(0.5,0.02,'Рис.5: Один слой Decoder (LLaMA-style) с GQA, SwiGLU, RMSNorm',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 6: Activation functions ───
def fig_activations():
    fig,ax=plt.subplots(figsize=(7,4)); sa(ax)
    x=np.linspace(-4,4,200)
    ax.plot(x,np.maximum(0,x),label='ReLU (Vaswani 2017)',color=C['blue'],linewidth=2)
    ax.plot(x,0.5*x*(1+np.tanh(np.sqrt(2/np.pi)*(x+0.044715*x**3))),':',label='GeLU (GPT-2/BERT)',color=C['green'],linewidth=2)
    ax.plot(x,x*(1/(1+np.exp(-x))),'--',label='SiLU (LLaMA 1)',color=C['purple'],linewidth=2)
    ax.plot(x,x*(1/(1+np.exp(-x)))*0.5,'-.',label='SwiGLU = SiLU·W₁ ⊙ W₂ (LLaMA 2/3)',color=C['orange'],linewidth=2)
    ax.set_xlabel('x',fontsize=9,color=C['text']); ax.set_ylabel('f(x)',fontsize=9,color=C['text'])
    ax.set_title('Активации FFN: от ReLU к SwiGLU',fontsize=10,color=C['text'],fontweight='bold')
    ax.legend(fontsize=7,facecolor=C['primary'],edgecolor=C['text_dim'],labelcolor=C['text'])
    ax.spines['left'].set_visible(True); ax.spines['bottom'].set_visible(True)
    fig.text(0.5,0.02,'Рис.6: Эволюция активационных функций FFN (Shazeer 2020, Touvron 2023)',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 7: Causal mask ───
def fig_causal_mask():
    fig,(ax1,ax2)=plt.subplots(1,2,figsize=(8,4))
    for ax in[ax1,ax2]: sa(ax)
    n=6
    unmasked=np.ones((n,n))*0.2; np.fill_diagonal(unmasked,0.8)
    masked=unmasked.copy()
    for i in range(n):
        for j in range(n):
            if j>i: masked[i,j]=0.001
    ax1.imshow(unmasked,cmap='viridis',aspect='auto',vmin=0,vmax=1)
    ax1.set_title('Без маски (неправильно)',fontsize=9,color=C['accent'],fontweight='bold')
    ax2.imshow(masked,cmap='viridis',aspect='auto',vmin=0,vmax=1)
    ax2.set_title('С causal mask (правильно)',fontsize=9,color=C['green'],fontweight='bold')
    for ax in[ax1,ax2]:
        ax.set_xticks(range(n)); ax.set_yticks(range(n))
        ax.set_xticklabels(range(1,n+1),fontsize=6,color=C['text'])
        ax.set_yticklabels(range(1,n+1),fontsize=6,color=C['text'])
    fig.text(0.5,0.02,'Рис.7: Causal Mask — почему модель не видит будущие токены',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 8: Generation ───
def fig_generation():
    fig,ax=plt.subplots(figsize=(9,4)); sa(ax); ax.axis('off')
    inp=FancyBboxPatch((0,0.55),2.8,0.35,boxstyle="round,pad=0.03",facecolor=C['green'],alpha=0.85,edgecolor='white')
    ax.add_patch(inp); ax.text(1.4,0.725,'ВХОД: "Какая наживка\nлучше для леща в..."',ha='center',va='center',fontsize=8,color='white',fontweight='bold')
    tr=FancyBboxPatch((3.3,0.5),1.8,0.45,boxstyle="round,pad=0.03",facecolor=C['secondary'],alpha=0.9,edgecolor=C['accent'],linewidth=2)
    ax.add_patch(tr); ax.text(4.2,0.725,'LLM\n(Decoder)',ha='center',va='center',fontsize=8,color='white',fontweight='bold')
    ax.annotate('',xy=(5.3,0.725),xytext=(5.1,0.725),arrowprops=dict(arrowstyle='->',color=C['text'],lw=2))
    probs={'июне':0.31,'августе':0.18,'июле':0.22,'мае':0.12,'сентябре':0.08,'апреле':0.05,'октябре':0.03,'ноябре':0.01}
    for i,(tok,p) in enumerate(probs.items()):
        x=5.6+i*0.4; h=p*4
        is_best=tok=='июне'
        cl=C['orange'] if is_best else C['blue']; a=0.8 if is_best else 0.3+p*2
        r=FancyBboxPatch((x-0.1,0.08),0.2,h,boxstyle="round,pad=0.01",facecolor=cl,alpha=a,edgecolor='white',linewidth=0.5 if not is_best else 1.5)
        ax.add_patch(r); ax.text(x,h+0.1,f'{p:.0%}',ha='center',va='bottom',fontsize=6,color=C['text'])
        ax.text(x,0.02,tok,ha='center',va='bottom',fontsize=6,color=C['text'])
    ax.text(5.6,0.7,'✓ июне\n(31%)',ha='center',fontsize=5.5,color=C['orange'],fontweight='bold')
    fig.text(0.5,0.02,'Рис.8: Авторегрессивная генерация — предсказание следующего токена с распределением вероятностей',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 9: Complexity ───
def fig_complexity():
    fig,(ax1,ax2)=plt.subplots(1,2,figsize=(9,4))
    for ax in[ax1,ax2]: sa(ax)
    n=np.array([128,256,512,1024,2048,4096]); d=768
    rnn_c=n*d**2; attn_c=n**2*d
    ax1.plot(n,rnn_c/1e6,'o-',label='RNN: O(n·d²)',color=C['blue'],lw=2)
    ax1.plot(n,attn_c/1e6,'s-',label='Transformer: O(n²·d)',color=C['accent'],lw=2)
    ax1.set_xlabel('seq len n',fontsize=8,color=C['text']); ax1.set_ylabel('MFLOPS',fontsize=8,color=C['text'])
    ax1.legend(fontsize=7,facecolor=C['primary'],edgecolor=C['text_dim'],labelcolor=C['text'])
    ax1.spines['left'].set_visible(True); ax1.spines['bottom'].set_visible(True)
    # memory
    mem=(n**2*2)/1e6 # FP16 bytes
    ax2.plot(n,mem,'d-',color=C['purple'],lw=2)
    ax2.set_xlabel('seq len n',fontsize=8,color=C['text']); ax2.set_ylabel('Attention matrix (MB)',fontsize=8,color=C['text'])
    ax2.axvline(x=2048,color=C['orange'],linestyle='--',alpha=0.5); ax2.text(2100,ax2.get_ylim()[1]*0.9,'наша модель',fontsize=7,color=C['orange'])
    ax2.spines['left'].set_visible(True); ax2.spines['bottom'].set_visible(True)
    fig.text(0.5,0.02,'Рис.9: Вычислительная сложность и потребление памяти Attention',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ─── FIG 10: Architecture comparison ───
def fig_arch_compare():
    fig,ax=plt.subplots(figsize=(8,4)); sa(ax); ax.axis('off')
    table_data=[['','Original TF','LLaMA 2/3','Наша модель'],
                ['Norm','LayerNorm','RMSNorm','RMSNorm'],
                ['PE','Sinusoidal','RoPE','RoPE'],
                ['FFN act','ReLU','SwiGLU','SwiGLU'],
                ['Attention','MHA','GQA','GQA (8 KV heads)'],
                ['Paras','65M-11B','7B-70B','~150M']]
    for i in range(6):
        for j in range(4):
            x=0.5+j*2.0; y=2.5-i*0.4
            cl=C['green'] if j==3 else C['blue'] if j==2 else C['secondary'] if j==1 else C['primary']
            if i==0: cl=C['orange']
            r=FancyBboxPatch((x-0.8,y-0.15),1.6,0.28,boxstyle="round,pad=0.02",facecolor=cl,alpha=0.5,edgecolor=cl,linewidth=0.5)
            ax.add_patch(r)
            ax.text(x,y,table_data[i][j],ha='center',va='center',fontsize=8 if i>0 else 9,color='white',fontweight='bold' if i==0 or j==3 else False)
    fig.text(0.5,0.02,'Рис.10: Сравнение архитектур — от оригинала Vaswani до LLaMA и нашей FishingLLM',ha='center',fontsize=8,color=C['text_dim'])
    return sf(fig)

# ═══ BUILD ═══
def build():
    doc=Document()
    style=doc.styles['Normal']; style.font.name='Calibri'; style.font.size=Pt(11)
    style.paragraph_format.space_after=Pt(6); style.paragraph_format.line_spacing=1.15

    # Title
    t=doc.add_heading('Лекция 1: Transformer — от теории до практической реализации',level=0); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=sub.add_run('FishingLLM — создание экспертной LLM по рыбалке в Тверской области\n'
                  'Источники: Vaswani et al. (2017), Su et al. (2021), Touvron et al. (2023),\n'
                  'Shazeer (2020), Kaplan et al. (2020), Hoffmann et al. (2022)')
    r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x66,0x66,0x66)
    doc.add_page_break()

    # ═══ ВВЕДЕНИЕ ═══
    doc.add_heading('Введение: зачем нам понимать Transformer',level=1)
    p=doc.add_paragraph()
    p.add_run('Наша цель — создать модель-эксперта по рыбалке в Тверской области. ')
    p.add_run('Для этого мы строим архитектуру Transformer с нуля — не потому что это модно, ')
    p.add_run('а потому что это единственная архитектура, которая:')
    doc.add_paragraph('• Обрабатывает все слова последовательности одновременно (параллельно)',style='List Bullet')
    doc.add_paragraph('• Масштабируется от 10M до 1T+ параметров',style='List Bullet')
    doc.add_paragraph('• Является стандартом индустрии (GPT, LLaMA, Mistral, Qwen)',style='List Bullet')
    doc.add_paragraph('• Имеет огромную экосистему инструментов (Hugging Face, llama.cpp, vLLM)',style='List Bullet')
    doc.add_picture(fig_history(),width=Inches(5.5))
    doc.add_page_break()

    # ═══ 1. ИСТОРИЯ ═══
    doc.add_heading('Глава 1. До Transformer: почему старые подходы не работали',level=1)
    doc.add_heading('1.1. RNN (Rumelhart et al., 1986)',level=2)
    p=doc.add_paragraph()
    p.add_run('Рекуррентные сети обрабатывают текст последовательно, слово за словом. '
              'Формально: hₜ = tanh(W·[hₜ₋₁; xₜ] + b). Проблема: ')
    r=p.add_run('затухающие градиенты')
    r.bold=True
    p.add_run('. При длине последовательности n>10 градиент умножается на Wª и '
              'экспоненциально затухает (или взрывается). Подробный анализ: Hochreiter (1991), Bengio et al. (1994).')

    doc.add_heading('1.2. LSTM (Hochreiter & Schmidhuber, 1997)',level=2)
    p=doc.add_paragraph()
    p.add_run('LSTM решила проблему затухания градиента через клеточное состояние Cₜ и шлюзы:')
    formulas=['fₜ = σ(Wf·[hₜ₋₁, xₜ] + bf) — забывающий шлюз',
              'iₜ = σ(Wi·[hₜ₋₁, xₜ] + bi) — входной шлюз',
              'Cₜ = fₜ·Cₜ₋₁ + iₜ·tanh(Wc·[hₜ₋₁, xₜ] + bc) — обновление памяти',
              'oₜ = σ(Wo·[hₜ₋₁, xₜ] + bo) — выходной шлюз',
              'hₜ = oₜ·tanh(Cₜ) — скрытое состояние']
    for f_text in formulas:
        p=doc.add_paragraph(f_text)
        if p.runs: p.runs[0].font.name='Consolas'; p.runs[0].font.size=Pt(9)

    p=doc.add_paragraph()
    p.add_run('НО: LSTM всё ещё последовательная — hₜ зависит от hₜ₋₁. '
              'Это означает O(n) шагов, без возможности параллелизации. '
              'На современных GPU это катастрофически медленно.')
    doc.add_page_break()

    # ═══ 2. ATTENTION ═══
    doc.add_heading('Глава 2. Внимание (Attention) — ключевая инновация',level=1)
    doc.add_heading('2.1. Откуда пришла идея',level=2)
    p=doc.add_paragraph()
    p.add_run('Bahdanau et al. (2015) впервые применили внимание в машинном переводе. '
              'Идея: на каждом шаге генерации "заглядывать" в исходное предложение и '
              'выбирать релевантные слова. Это называлось additive attention.')
    p=doc.add_paragraph()
    p.add_run('Vaswani et al. (2017) предложили ')
    r=p.add_run('dot-product attention')
    r.bold=True
    p.add_run(' — более эффективную версию, которая стала стандартом. '
              'Основное преимущество: возможность использовать оптимизированные '
              'матричные умножения (GEMM) на GPU.')

    doc.add_heading('2.2. Query, Key, Value — аналогия',level=2)
    p=doc.add_paragraph()
    p.add_run('Представьте поиск в Google:')
    p=doc.add_paragraph('• Q (Query) = ваш поисковый запрос: "ловля щуки на спиннинг в Тверской области"')
    p=doc.add_paragraph('• K (Key) = заголовки страниц: "Ловля щуки", "Рыбалка в Тверской области"')
    p=doc.add_paragraph('• V (Value) = содержимое этих страниц')
    p=doc.add_paragraph()
    p.add_run('Google берёт ваш запрос (Q), сравнивает со всеми заголовками (K), '
              'оценивает релевантность и возвращает содержимое (V) самых подходящих страниц. '
              'Self-Attention делает то же самое — но на уровне токентов внутри одного предложения.')

    doc.add_heading('2.3. Формула Attention — посимвольный разбор',level=2)
    p=doc.add_paragraph()
    r=p.add_run('Attention(Q, K, V) = softmax( Q·Kᵀ / √dₖ ) · V')
    r.font.size=Pt(13); r.font.name='Consolas'; r.bold=True
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    items=[
        ('Q [seq_len × dₖ]','Матрица запросов. Каждый токен "спрашивает": кто здесь важен для меня?'),
        ('Kᵀ [dₖ × seq_len]','Транспонированная матрица ключей. Каждый токен "отвечает": я важен настолько-то.'),
        ('Q·Kᵀ [seq_len × seq_len]','Матрица внимания. Элемент (i,j) — скалярное произведение запроса i и ключа j. Чем больше → тем сильнее токен i "обращает внимание" на токен j.'),
        ('√dₖ','Масштабирование. dₖ — размерность ключа (обычно 64-128). Без него при большой dₖ значения становятся слишком большими, softmax превращается в argmax, градиент исчезает.'),
        ('softmax','Нормализация по строкам. Каждая строка суммируется в 1. Это создаёт "соревнование" между токенами за внимание.'),
        ('V [seq_len × dᵥ]','Матрица значений. Информация, которую токены передают друг другу.'),
        ('·V','Взвешенная сумма. Каждый токен получает смесь информации от всех остальных, пропорционально весам внимания.'),
    ]
    for sym,desc in items:
        p=doc.add_paragraph()
        r=p.add_run(f'{sym}: '); r.bold=True; r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x3B,0x82,0xF6)
        p.add_run(desc)

    doc.add_picture(fig_qkv(),width=Inches(5.5))
    doc.add_page_break()

    doc.add_heading('2.4. Числовой пример: Self-Attention на словах "щука лещ"',level=2)
    p=doc.add_paragraph()
    p.add_run('Разберём конкретный пример. Пусть у нас два токена с эмбеддингами размерности d=3:')
    p=doc.add_paragraph()
    p.add_run('"щука" → [0.8, 0.2, 0.5],  "лещ" → [0.1, 0.9, 0.3]')
    p=doc.add_paragraph()
    p.add_run('Шаг 1: Q·Kᵀ — матрица 2×2:')
    p=doc.add_paragraph('  [0.93, 0.49]   # щука·щука=0.93, щука·лещ=0.49')
    p=doc.add_paragraph('  [0.49, 0.91]   # лещ·щука=0.49, лещ·лещ=0.91')
    p=doc.add_paragraph()
    p.add_run('Шаг 2: делим на √3≈1.73, softmax:')
    p=doc.add_paragraph('  [0.56, 0.44]   # щука: 56% себе, 44% лещу')
    p=doc.add_paragraph('  [0.44, 0.56]   # лещ: 44% щуке, 56% себе')
    p=doc.add_paragraph()
    p.add_run('Шаг 3: взвешенная сумма с V:')
    p=doc.add_paragraph('  "щука" → 0.56·[0.8,0.2,0.5] + 0.44·[0.1,0.9,0.3] = [0.49, 0.51, 0.41]')
    p=doc.add_paragraph('  "лещ"  → 0.44·[0.8,0.2,0.5] + 0.56·[0.1,0.9,0.3] = [0.41, 0.59, 0.39]')
    p=doc.add_paragraph()
    p.add_run('Интерпретация: каждый токен "позаимствовал" ~40-45% информации от соседа. '
              'Эмбеддинги стали более похожи — модель поняла, что щука и лещ — родственные концепции (рыбы). ')
    r=p.add_run('Это и есть контекстуализация.')
    r.bold=True

    doc.add_picture(fig_attention_numerical(),width=Inches(5.5))
    doc.add_page_break()

    doc.add_heading('2.5. Маскированное внимание (Causal Mask)',level=2)
    p=doc.add_paragraph()
    p.add_run('При генерации модель не должна "подглядывать" в будущие токены. '
              'Маска — матрица [seq_len × seq_len], где элемент (i,j):')
    p=doc.add_paragraph('  = 0, если j ≤ i (можно смотреть на текущий и предыдущие токены)')
    p=doc.add_paragraph('  = -∞, если j > i (нельзя — это будущее)')
    p=doc.add_paragraph()
    code=doc.add_paragraph()
    r=code.add_run(
        '# Маска для seq_len=4:\n'
        'mask = [[  0, -∞, -∞, -∞],\n'
        '        [  0,  0, -∞, -∞],\n'
        '        [  0,  0,  0, -∞],\n'
        '        [  0,  0,  0,  0]]\n\n'
        '# После softmax всё, что было -∞, стало 0:\n'
        'attn = [[0.8, 0.0, 0.0, 0.0],  # токен 1 видит только себя\n'
        '        [0.4, 0.6, 0.0, 0.0],  # токен 2 видит 1 и 2\n'
        '        ...]'
    )
    r.font.name='Consolas'; r.font.size=Pt(8)
    doc.add_picture(fig_causal_mask(),width=Inches(5))

    p=doc.add_paragraph()
    p.add_run('Важно: ')
    r=p.add_run('маска применяется ДО softmax. -∞ после softmax даёт ровно 0. '
                'Модель физически не может получить информацию из будущего.')
    r.bold=True
    doc.add_page_break()

    doc.add_heading('2.6. Почему dot-product, а не additive?',level=2)
    p=doc.add_paragraph()
    p.add_run('Vaswani et al. экспериментально сравнили оба подхода:')
    p=doc.add_paragraph('• Additive attention (Bahdanau 2015): более выразительна, но медленнее ×2-3',style='List Bullet')
    p=doc.add_paragraph('• Dot-product attention: быстрее (оптимизированные BLAS-операции), качество сопоставимо',style='List Bullet')
    p=doc.add_paragraph()
    p.add_run('Выбор: dot-product. Для нашей задачи он даёт достаточную выразительность при минимальных вычислительных затратах.')
    doc.add_page_break()

    # ═══ 3. MULTI-HEAD ═══
    doc.add_heading('Глава 3. Multi-Head Attention — несколько точек зрения',level=1)
    p=doc.add_paragraph()
    p.add_run('Одна голова внимания — это один "взгляд" на текст. Но предложение '
              '"Лещ клюёт на червя в июне на Верхней Волге" содержит разные типы связей:')
    p=doc.add_paragraph('• Синтаксические: подлежащее-сказуемое (Лещ — клюёт)',style='List Bullet')
    p=doc.add_paragraph('• Семантические: объект-действие (червя — клюёт)',style='List Bullet')
    p=doc.add_paragraph('• Временные: месяц-событие (июнь — клюёт)',style='List Bullet')
    p=doc.add_paragraph('• Пространственные: локация-действие (Волга — клюёт)',style='List Bullet')
    p=doc.add_paragraph()
    p.add_run('Multi-Head Attention = h параллельных голов, каждая со своими Qᵢ, Kᵢ, Vᵢ. '
              'Формула: MultiHead(Q,K,V) = Concat(head₁,...,headₕ)·Wᵒ.')
    p=doc.add_paragraph()
    p.add_run('Сравнение количества голов:')
    table=doc.add_table(rows=5,cols=4); table.style='Light Shading Accent 1'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    cells=[['Модель','Параметры','n_heads','d_k'],
           ['Original TF','65M','8','64'],
           ['LLaMA 7B','7B','32','128'],
           ['GPT-3','175B','96','128'],
           ['Наша FishingLLM','~150M','12','64']]
    for i,row in enumerate(cells):
        for j,val in enumerate(row):
            table.cell(i,j).text=val
            if i==0:
                for pr in table.cell(i,j).paragraphs:
                    for rr in pr.runs: rr.bold=True
            if i==4:
                for pr in table.cell(i,j).paragraphs:
                    for rr in pr.runs: rr.bold=True; rr.font.color.rgb=RGBColor(0xE9,0x45,0x60)
    doc.add_page_break()

    # ═══ 4. POSITIONAL ENCODING ═══
    doc.add_heading('Глава 4. Позиционное кодирование: как модель узнаёт порядок слов',level=1)
    p=doc.add_paragraph()
    p.add_run('Transformer обрабатывает все слова параллельно → теряет информацию о порядке. '
              '"Щука поймала рыбака" и "Рыбака поймала щука" — для модели одинаковый набор токенов.')

    doc.add_heading('4.1. Синусоидальное кодирование (Vaswani 2017)',level=2)
    p=doc.add_paragraph()
    p.add_run('PE(pos, 2i) = sin(pos / 10000^{2i/d})')
    p=doc.add_paragraph()
    p.add_run('PE(pos, 2i+1) = cos(pos / 10000^{2i/d})')
    p=doc.add_paragraph()
    p.add_run('Каждая размерность — синусоида со своей частотой. Низкие размерности '
              'меняются быстро (кодируют близкие позиции), высокие — медленно (глобальное положение).')
    p=doc.add_paragraph()
    r=p.add_run('Недостаток: ')
    r.bold=True
    p.add_run('фиксированная длина. После обучения на n=1024, модель не может '
              'обобщить на n=2048 без дообучения.')

    doc.add_heading('4.2. RoPE — Rotary Position Embedding (Su et al., 2021)',level=2)
    p=doc.add_paragraph()
    p.add_run('Вместо добавления сигнала к эмбеддингу, RoPE ')
    r=p.add_run('поворачивает векторы Q и K')
    r.bold=True
    p.add_run(' в комплексной плоскости в зависимости от их позиции.')

    p=doc.add_paragraph()
    p.add_run('Формально: f(qₘ, m) = qₘ · e^(i·m·θ), где θₖ = 1/10000^{2k/d_model}')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

    p=doc.add_paragraph()
    p.add_run('Ключевое свойство: ')
    r=p.add_run('⟨Qₘ, Kₙ⟩ = ⟨Q₀, Kₙ₋ₘ⟩')
    r.bold=True; r.font.name='Consolas'; r.font.size=Pt(10)
    p.add_run(' — скалярное произведение зависит ТОЛЬКО от относительного расстояния '
              'между токенами (|m-n|), а не от их абсолютных позиций.')

    p=doc.add_paragraph()
    p.add_run('Преимущества RoPE:')
    p=doc.add_paragraph('• Относительное позиционирование — модель понимает "на сколько слов назад"',style='List Bullet')
    p=doc.add_paragraph('• Лучшая экстраполяция — можно увеличить контекст после обучения',style='List Bullet')
    p=doc.add_paragraph('• Естественная decay — внимание к дальним токенам уменьшается',style='List Bullet')
    p=doc.add_paragraph('• Стандарт в LLaMA, Mistral, Qwen — поддерживается всеми фреймворками',style='List Bullet')

    p=doc.add_paragraph()
    r=p.add_run('Для FishingLLM: используем RoPE с base=10000, partial_factor=0.25 — '
                '25% размерностей участвуют во вращении (LLaMA-style).')
    r.bold=True

    doc.add_picture(fig_rope_full(),width=Inches(5.5))
    doc.add_page_break()

    # ═══ 5. FULL LAYER ═══
    doc.add_heading('Глава 5. Полный слой декодера — сборка компонентов',level=1)
    p=doc.add_paragraph('Один слой Transformer-декодера содержит:')
    p=doc.add_paragraph('1. Вход: X + Positional Encoding (или RoPE применяется к Q/K)')
    p=doc.add_paragraph('2. RMS Norm — стабилизация распределения активаций')
    p=doc.add_paragraph('3. Grouped Query Attention — замена MHA для эффективности')
    p=doc.add_paragraph('4. + Residual connection')
    p=doc.add_paragraph('5. RMS Norm — вторая нормализация')
    p=doc.add_paragraph('6. SwiGLU FFN — нелинейное преобразование')
    p=doc.add_paragraph('7. + Residual connection')

    doc.add_picture(fig_decoder(),width=Inches(4.5))

    doc.add_heading('5.1. RMS Norm — замена LayerNorm',level=2)
    p=doc.add_paragraph()
    p.add_run('LayerNorm: LN(x) = (x - μ)/√(σ² + ε)·γ + β. Центрирование (μ) требует дополнительных вычислений.')
    p=doc.add_paragraph()
    p.add_run('RMSNorm (Zhang & Sennrich, 2019): ')
    r=p.add_run('RMSNorm(x) = x / √(mean(x²) + ε)·γ')
    r.bold=True; r.font.name='Consolas'; r.font.size=Pt(10)
    p.add_run('. Без центрирования! На ~15% быстрее, качество сопоставимо. '
              'γ и β — обучаемые параметры размерности d_model. Без них сеть '
              'не могла бы сдвигать/масштабировать распределения под свои нужды.')

    doc.add_heading('5.2. FFN и эволюция активаций',level=2)
    p=doc.add_paragraph()
    p.add_run('FFN — это "кабинет для размышлений": после внимания (собрания) '
              'каждый токен обрабатывает полученную информацию.')

    p=doc.add_paragraph()
    p.add_run('Эволюция:')
    p=doc.add_paragraph('• ReLU (Vaswani 2017): f(x)=max(0,x). Просто, но "умирающие нейроны" при x<0.',style='List Bullet')
    p=doc.add_paragraph('• GeLU (GPT-2, BERT): гладкая аппроксимация ReLU, f(x)=x·Φ(x). Качество выше.',style='List Bullet')
    p=doc.add_paragraph('• SiLU (LLaMA 1): f(x)=x·σ(x). Ещё глаже, ε=0.',style='List Bullet')
    p=doc.add_paragraph('• SwiGLU (LLaMA 2/3, Mistral, Qwen): SiLU(x·W₁) ⊙ (x·W₂) · W₃. '
                         'Лучшее качество, но +50% параметров (3 матрицы вместо 2).',style='List Bullet')

    p=doc.add_paragraph()
    r=p.add_run('Для FishingLLM: SwiGLU с d_ff=4·d_model.')
    r.bold=True
    p.add_run(' Формула: FFN_SwiGLU(x) = (SiLU(x·W₁) ⊙ (x·W₂)) · W₃')

    doc.add_picture(fig_activations(),width=Inches(5))
    doc.add_page_break()

    doc.add_heading('5.3. Grouped Query Attention (GQA)',level=2)
    p=doc.add_paragraph()
    p.add_run('В оригинальном Transformer — Multi-Head Attention (MHA): n_query = n_key = n_value = n_heads.')
    p=doc.add_paragraph()
    p.add_run('Проблема MHA: каждая KV-голова требует памяти под Key/Value кэш. '
              'Для генерации (autoregressive) это O(n·d·n_heads) на каждый запрос.')
    p=doc.add_paragraph()
    p.add_run('Решение — GQA (Ainslie et al., 2023, LLaMA 2): ')
    r=p.add_run('группируем Query-головы, чтобы они делили общие KV-головы.')
    r.bold=True

    p=doc.add_paragraph()
    p.add_run('Пример: n_heads=12, n_kv_heads=4. Каждая KV-голова обслуживает 3 Query-головы. '
              'KV-кэш уменьшается в 3 раза, качество падает незначительно.')
    p=doc.add_paragraph()
    r=p.add_run('Для FishingLLM: n_heads=12, n_kv_heads=4.')
    r.bold=True
    doc.add_page_break()

    # ═══ 6. TRAINING ═══
    doc.add_heading('Глава 6. Обучение и генерация',level=1)

    doc.add_heading('6.1. Языковое моделирование (Next Token Prediction)',level=2)
    p=doc.add_paragraph()
    p.add_run('LLM обучается предсказывать следующий токен:')
    p=doc.add_paragraph()
    r=p.add_run('P(xₜ | x₁, ..., xₜ₋₁) = softmax(hₜ · W_out)')
    r.font.name='Consolas'; r.font.size=Pt(10)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph()
    p.add_run('Функция потерь — кросс-энтропия:')
    r=p.add_run('L = -1/N · Σ log P(xₜ | x_{<t})')
    r.font.name='Consolas'; r.font.size=Pt(10)
    p=doc.add_paragraph()
    p.add_run('Интуиция: loss — это логарифм вероятности, которую модель '
              'присвоила правильному токену. Если loss=0 — модель "уверена" '
              'на 100%. На практике loss обычно 2-4 (для BPE vocab 32K).')

    doc.add_heading('6.2. Стратегии генерации',level=2)
    p=doc.add_paragraph()
    p.add_run('Greedy: argmax — всегда самый вероятный токен. '
              'Просто, но склонно к повторам.')
    p=doc.add_paragraph()
    p.add_run('Top-k: выбрать k самых вероятных (k=40), ресемплировать.')
    p=doc.add_paragraph()
    p.add_run('Top-p (nucleus): выбрать минимальное множество с кумулятивной '
              'вероятностью p (p=0.9). Адаптивно.')
    p=doc.add_paragraph()
    p.add_run('Temperature: масштабировать логиты. T<1 — острее, T>1 — площе.')
    p=doc.add_paragraph()
    r=p.add_run('Для FishingLLM: top-p=0.9, temperature=0.7.')
    r.bold=True

    doc.add_picture(fig_generation(),width=Inches(5.5))
    doc.add_page_break()

    # ═══ 7. BACKPROP ═══
    doc.add_heading('Глава 7. Обратное распространение через Attention',level=1)
    p=doc.add_paragraph()
    p.add_run('Как градиент (сигнал ошибки) проходит через Attention?')
    p=doc.add_paragraph()
    p.add_run('Цепочка: Output ← W←V ← softmax ← S_scaled ← S_raw ← Q, K')
    p=doc.add_paragraph()

    p=doc.add_paragraph('Шаг 1: ∂L/∂V = Wᵀ · ∂L/∂Output',style='List Number')
    p=doc.add_paragraph('  — градиент от выхода к V через веса внимания')
    p=doc.add_paragraph('Шаг 2: ∂L/∂S_scaled = ∂L/∂Output · Vᵀ',style='List Number')
    p=doc.add_paragraph('  — градиент от V к scores (обратное направление умножения)')
    p=doc.add_paragraph('Шаг 3: ∂L/∂S_raw = ∂L/∂S_scaled · 1/√dₖ',style='List Number')
    p=doc.add_paragraph('  — масштабирование')
    p=doc.add_paragraph('Шаг 4: ∂L/∂S_raw[i] через softmax',style='List Number')
    p2=doc.add_paragraph('  ∂L/∂S_raw[i] = S[i] · (∂L/∂S_scaled[i] - Σ(S[j] · ∂L/∂S_scaled[j]))')
    r=p2.runs[0]
    r.font.name='Consolas'; r.font.size=Pt(9)
    p=doc.add_paragraph('  — самый важный шаг: из-за нормализации softmax градиент '
                         'к одному токену зависит от всех других!')
    p=doc.add_paragraph('Шаг 5: ∂L/∂Q = ∂L/∂S_raw · K, ∂L/∂K = Qᵀ · ∂L/∂S_raw',style='List Number')
    p=doc.add_paragraph('Шаг 6: ∂L/∂W_Q = Xᵀ · ∂L/∂Q, и т.д.',style='List Number')

    p=doc.add_paragraph()
    r=p.add_run('Важное следствие: ')
    r.bold=True
    p.add_run('из-за softmax градиенты к разным токенам взаимозависимы. '
              'Если один токен получает вес 0.9, другой получит 0.1, '
              'и его градиент будет подавлен. Это создаёт "соревнование" '
              'за внимание, что делает обучение динамичным.')

    p=doc.add_paragraph()
    p.add_run('Flash Attention (Dao et al., 2022): ')
    r=p.add_run('оптимизация, которая вычисляет Attention блоками, '
                'не сохраняя матрицу [n×n] целиком. Ускорение ×2-4, '
                'снижение памяти с O(n²) до O(n). Будем использовать при обучении.')
    r.bold=True
    doc.add_page_break()

    # ═══ 8. COMPLEXITY ═══
    doc.add_heading('Глава 8. Вычислительная сложность',level=1)
    p=doc.add_paragraph()
    p.add_run('Основные формулы:')
    p=doc.add_paragraph('  Attention: O(n²·d) — матрица [n×n]')
    p=doc.add_paragraph('  FFN: O(n·d²) — умножения на W₁, W₂')
    p=doc.add_paragraph('  Итого на слой: O(n²·d + n·d²)')

    p=doc.add_paragraph()
    p.add_run('Для FishingLLM (n=1024, d=768, L=12):')
    p=doc.add_paragraph('  Attention: 1024²·768 ≈ 805 MFLOPS/слой')
    p=doc.add_paragraph('  FFN (SwiGLU): 1024·(768·3072·2) ≈ 4.8 GFLOPS/слой')
    p=doc.add_paragraph('  Всего: 12 слоёв × 5.6 GFLOPS ≈ 67 GFLOPS на forward pass')
    p=doc.add_paragraph('  Для обучения (forward + backward): ×3 ≈ 200 GFLOPS/шаг')
    p=doc.add_paragraph('  На T4 (8.1 TFLOPS FP16): ~40 шагов/сек')

    doc.add_picture(fig_complexity(),width=Inches(5.5))
    p=doc.add_paragraph()
    p.add_run('Вывод: ')
    r=p.add_run('для нашего контекста (n=1024) доминирует FFN, а не Attention. '
                'O(n²·d) становится проблемой при n > 4096. Для нас — не критично.')
    r.bold=True
    doc.add_page_break()

    # ═══ 9. WHY TRANSFORMER WON ═══
    doc.add_heading('Глава 9. Почему Transformer "победил"',level=1)
    p=doc.add_paragraph()
    p=doc.add_paragraph('1. Параллелизация: O(1) шагов против O(n) у RNN')
    p=doc.add_paragraph('2. Дальние зависимости: любой токен видит любой за 1 шаг')
    p=doc.add_paragraph('3. Scaling Laws (Kaplan et al., 2020): качество предсказуемо '
                         'растёт с увеличением N_params, D_data и FLOPs')
    p=doc.add_paragraph('4. Chinchilla (Hoffmann et al., 2022): оптимальное соотношение '
                         'параметров и данных — 20 токенов на 1 параметр')
    p=doc.add_paragraph('5. Индуктивные bias: нет сильных предубеждений → учит любые паттерны')
    p=doc.add_paragraph('6. Экосистема: Hugging Face, TensorRT, llama.cpp, vLLM, Flash Attention')
    doc.add_page_break()

    # ═══ 10. MODERN ARCHITECTURE ═══
    doc.add_heading('Глава 10. Современная архитектура LLM (2024-2025)',level=1)
    p=doc.add_paragraph()
    p.add_run('Все современные LLM используют один и тот же шаблон:')
    p=doc.add_paragraph('• Decoder-only (не Encoder-Decoder) — для генерации')
    p=doc.add_paragraph('• Pre-RMSNorm (нормализация ДО Attention/FFN, а не после)')
    p=doc.add_paragraph('• RoPE (вместо синусоид)')
    p=doc.add_paragraph('• SwiGLU (вместо ReLU)')
    p=doc.add_paragraph('• GQA (вместо MHA, для моделей >1B)')
    p=doc.add_paragraph('• Flash Attention (для эффективного обучения)')

    doc.add_picture(fig_arch_compare(),width=Inches(5.5))

    p=doc.add_paragraph()
    p.add_run('Альтернатива — SSM / Mamba (Gu & Dao, 2023): ')
    r=p.add_run('сложность O(n) вместо O(n²), контекст до 1M+ токенов. '
                'НО: хуже на задачах фактологического запоминания, '
                'меньше инструментов, сложнее в обучении.')
    r.bold=True
    p.add_run(' Для FishingLLM выбираем Transformer — контекст небольшой '
              '(2048), а качество и экосистема важнее.')
    doc.add_page_break()

    # ═══ 11. OUR MODEL ═══
    doc.add_heading('Глава 11. Проект FishingLLM-150M — архитектурные решения',level=1)
    p=doc.add_paragraph()
    p.add_run('Сводка решений с обоснованием:')

    table=doc.add_table(rows=9,cols=3); table.style='Light Shading Accent 1'
    headers=['Параметр','Значение','Обоснование']
    for j,h in enumerate(headers): table.cell(0,j).text=h
    rows=[['vocab_size','16 000','BPE-токенизатор под русский корпус (компактнее 32K)'],
          ['d_model','768','Стандарт для ~150M. Позволяет d_k=64 при n_heads=12'],
          ['n_layers','12','LLaMA-подобное соотношение глубина/ширина'],
          ['n_heads','12','d_model / d_k = 768/64 = 12'],
          ['n_kv_heads','4','GQA: каждая KV-голова обслуживает 3 Q-головы'],
          ['d_ff','3072','SwiGLU: 3·d_model·d_ff ≈ 3·768·3072 = 7.08M/слой'],
          ['max_seq_len','2048','Достаточно для статей и диалогов о рыбалке'],
          ['total_params','~150M','Embed: 12.3M + 12 слоёв × 10M + LM head: 12.3M']]
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            table.cell(i+1,j).text=val
            if i==len(rows)-1:
                for pr in table.cell(i+1,j).paragraphs:
                    for rr in pr.runs: rr.bold=True; rr.font.color.rgb=RGBColor(0xE9,0x45,0x60)

    doc.add_paragraph('')
    p=doc.add_paragraph()
    r=p.add_run('Scaling Law для FishingLLM: ')
    r.bold=True
    p.add_run('по правилу Chinchilla (20 токенов на 1 параметр), '
              '150M·20 = 3B токенов. Наш корпус на русском языке '
              'должен быть не менее ~3 млрд токенов для полноценного обучения.')
    doc.add_page_break()

    # ═══ 12. PRACTICE ═══
    doc.add_heading('Глава 12. Практические задания для закрепления',level=1)

    p=doc.add_paragraph()
    r=p.add_run('Задание 1. Реализуйте Self-Attention на PyTorch')
    r.bold=True
    code=doc.add_paragraph()
    r=code.add_run(
        'import torch\n'
        'import torch.nn.functional as F\n\n'
        'class SelfAttention(torch.nn.Module):\n'
        '    def __init__(self, d_model, n_heads):\n'
        '        super().__init__()\n'
        '        self.d_k = d_model // n_heads\n'
        '        self.W_Q = torch.nn.Linear(d_model, d_model)\n'
        '        self.W_K = torch.nn.Linear(d_model, d_model)\n'
        '        self.W_V = torch.nn.Linear(d_model, d_model)\n'
        '        self.W_O = torch.nn.Linear(d_model, d_model)\n\n'
        '    def forward(self, x, mask=None):\n'
        '        # x: [batch, seq, d_model]\n'
        '        Q = self.W_Q(x)\n'
        '        K = self.W_K(x)\n'
        '        V = self.W_V(x)\n'
        '        scores = Q @ K.transpose(-2,-1) / (self.d_k**0.5)\n'
        '        if mask is not None:\n'
        '            scores = scores + mask\n'
        '        weights = F.softmax(scores, dim=-1)\n'
        '        return self.W_O(weights @ V)\n\n'
        '# Проверка на синтетических данных:\n'
        'x = torch.randn(2, 8, 768)  # batch=2, seq=8, d=768\n'
        'attn = SelfAttention(768, 12)\n'
        'out = attn(x)\n'
        'print(out.shape)  # [2, 8, 768] — размерность не изменилась'
    )
    r.font.name='Consolas'; r.font.size=Pt(8)

    p=doc.add_paragraph()
    r=p.add_run('Задание 2. Визуализируйте матрицу внимания')
    r.bold=True
    p=doc.add_paragraph('Для фразы "На какую наживку клюёт лещ в июне на Верхней Волге" '
                         'постройте heatmap матрицы внимания. Определите топ-3 самых '
                         'сильных связей между токенами. Совпадает ли это с '
                         'вашей лингвистической интуицией?')

    p=doc.add_paragraph()
    r=p.add_run('Задание 3. Сравните время выполнения Attention')
    r.bold=True
    p=doc.add_paragraph('Напишите бенчмарк для seq_len = [128, 256, 512, 1024, 2048]. '
                         'Постройте график времени выполнения. '
                         'При какой длине время инференса превышает 100 мс?')

    p=doc.add_paragraph()
    r=p.add_run('Задание 4 (research). Изучите конфиги реальных моделей')
    r.bold=True
    p=doc.add_paragraph('На Hugging Face найдите config.json для: '
                         'TinyLlama-1.1B, Qwen2.5-1.5B, GPT-2. '
                         'Сравните их архитектурные решения с нашей FishingLLM. '
                         'Какие отличия вы видите и чем они могут быть обоснованы?')

    p=doc.add_paragraph()
    r=p.add_run('Задание 5. Подсчитайте параметры модели')
    r.bold=True
    p=doc.add_paragraph('Используя конфиг из Главы 11, вручную подсчитайте '
                         'общее число параметров. Проверьте: сходится ли с ~150M?')

    doc.add_paragraph('')

    p=doc.add_paragraph()
    r=p.add_run('Критерии проверки:')
    r.bold=True
    doc.add_paragraph('• Задание 1: выход не меняет размерность [b,n,d]',style='List Bullet')
    doc.add_paragraph('• Задание 2: max weight > 0.5 — на диагонали или на тематической паре',style='List Bullet')
    doc.add_paragraph('• Задание 3: график показывает квадратичный рост',style='List Bullet')
    doc.add_paragraph('• Задание 5: результат ±5% от 150M',style='List Bullet')

    doc.add_paragraph('')
    p=doc.add_paragraph()
    r=p.add_run('Литература:')
    r.bold=True; r.font.size=Pt(12)
    refs=['Vaswani et al. (2017) — "Attention Is All You Need". arXiv:1706.03762',
          'Su et al. (2021) — "RoFormer: Enhanced Transformer with Rotary Position Embedding". arXiv:2104.09864',
          'Touvron et al. (2023) — "LLaMA: Open and Efficient Foundation Language Models". arXiv:2302.13971',
          'Shazeer (2020) — "GLU Variants Improve Transformer". arXiv:2002.05202',
          'Ainslie et al. (2023) — "GQA: Training Generalized Multi-Query Transformer Models". arXiv:2305.13245',
          'Dao et al. (2022) — "Flash Attention: Fast and Memory-Efficient Exact Attention". arXiv:2205.14135',
          'Kaplan et al. (2020) — "Scaling Laws for Neural Language Models". arXiv:2001.08361',
          'Hoffmann et al. (2022) — "Training Compute-Optimal Large Language Models". arXiv:2203.15556',
          'Zhang & Sennrich (2019) — "Root Mean Square Layer Normalization". arXiv:1910.07467',
          '3Blue1Brown (2024) — "Attention in transformers, visually explained". YouTube',
          'Karpathy (2023) — "Let\'s build GPT from scratch". GitHub (nanoGPT)']
    for ref in refs:
        p=doc.add_paragraph(ref)
        p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=RGBColor(0x66,0x66,0x66)

    doc.add_paragraph('')
    p=doc.add_paragraph()
    r=p.add_run('Следующий шаг: Лекция 2 — Tokenization. '
                'Научимся превращать текст на русском языке в числа.')
    r.bold=True; r.font.color.rgb=RGBColor(0x0f,0x34,0x60)

    os.makedirs(OUTPUT, exist_ok=True)
    path=os.path.join(OUTPUT,'lecture-01-transformer.docx')
    doc.save(path)
    print(f'Документ сохранён: {path}')

if __name__=='__main__':
    build()
