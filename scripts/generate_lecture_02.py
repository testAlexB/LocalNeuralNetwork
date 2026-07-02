"""Генерация Word-документа: Лекция 2 — Tokenization (теория + практика)"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io, os, sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')

COLORS = {
    'bg': '#1a1a2e', 'primary': '#16213e', 'secondary': '#0f3460',
    'accent': '#e94560', 'orange': '#f59e0b', 'green': '#10b981',
    'blue': '#3b82f6', 'purple': '#8b5cf6', 'teal': '#06b6d4',
    'text': '#e2e8f0', 'text_dim': '#94a3b8',
}

def setup_style(ax):
    ax.set_facecolor(COLORS['bg'])
    ax.figure.patch.set_facecolor(COLORS['bg'])
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(colors=COLORS['text_dim'], labelsize=6)
    ax.grid(True, alpha=0.08, color=COLORS['text_dim'])

def save_fig(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor=COLORS['bg'], edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


# ============================================================
# ДИАГРАММА 1: Как работает BPE — пошагово
# ============================================================
def diagram_bpe_steps():
    fig, ax = plt.subplots(figsize=(8, 5))
    setup_style(ax)
    ax.set_xlim(-1, 9)
    ax.set_ylim(0, 3.5)
    ax.axis('off')

    steps = [
        (0, 'Шаг 0\nИсходный текст', '"щука лещ окунь"\n→ разбиваем на символы'),
        (2, 'Шаг 1\nСчёт пар', '("щ","у"): 1\n("у","к"): 1\n("к","а"): 1\n...'),
        (4, 'Шаг 2\nСлияние', 'Самая частая пара\n"л"+"е" → "ле"\nСловарь: +1 токен'),
        (6, 'Шаг N\nРезультат', 'Словарь готов:\n"щук", "а", "ле", "щ",\n"_" и т.д.'),
    ]

    for x, title, content in steps:
        rect = FancyBboxPatch((x, 2.2), 1.8, 0.8,
                              boxstyle="round,pad=0.05",
                              facecolor=COLORS['secondary'], alpha=0.85,
                              edgecolor=COLORS['blue'], linewidth=1)
        ax.add_patch(rect)
        ax.text(x + 0.9, 2.6, title, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold')
        ax.text(x + 0.9, 1.5, content, ha='center', va='top',
                fontsize=7, color=COLORS['text'])

        if x > 0:
            ax.annotate('', xy=(x, 2.6), xytext=(x - 0.2, 2.6),
                        arrowprops=dict(arrowstyle='->', color=COLORS['accent'], lw=2))

    # Снизу — общий принцип
    ax.text(4, 0.5, 'BPE iteratively merges the most frequent pair of adjacent tokens\n'
                    'BPE итеративно сливает самую частую пару соседних токенов',
            ha='center', fontsize=9, color=COLORS['orange'])

    fig.text(0.5, 0.02, 'Рис. 1: Пошаговый процесс Byte Pair Encoding',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 2: Сравнение токенизаторов
# ============================================================
def diagram_tokenizer_comparison():
    fig, ax = plt.subplots(figsize=(8, 4))
    setup_style(ax)
    ax.set_xlim(-1, 8)
    ax.set_ylim(0, 2.5)
    ax.axis('off')

    data = [
        ('Word-level\n(пословно)', '["щука", "лещ", "окунь",\n"поймал", ...]',
         'Словарь: ∞\nРедкие слова\nнеизвестны', COLORS['accent']),
        ('Char-level\n(посимвольно)', '["щ", "у", "к", "а",\n"_", "л", "е", "щ", ...]',
         'Словарь: ~100\nДлинные\nпоследовательности', COLORS['purple']),
        ('BPE\n(подсловно)', '["щук", "а", "_", "ле",\n"щ", "_", "окунь"]',
         'Словарь: 32K–128K\nБаланс размера\nи покрытия ✓', COLORS['green']),
    ]

    for i, (title, example, desc, color) in enumerate(data):
        x = i * 2.6
        rect = FancyBboxPatch((x, 1.3), 2.2, 0.9,
                              boxstyle="round,pad=0.03",
                              facecolor=color, alpha=0.15,
                              edgecolor=color, linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + 1.1, 1.75, title, ha='center', va='center',
                fontsize=9, color='white', fontweight='bold')
        ax.text(x + 1.1, 1.1, example, ha='center', va='top',
                fontsize=7, color=COLORS['text'])
        ax.text(x + 1.1, 0.35, desc, ha='center', va='top',
                fontsize=7, color=color)

    fig.text(0.5, 0.02, 'Рис. 2: Сравнение подходов к токенизации',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 3: Русская рыбалка — визуализация токенов
# ============================================================
def diagram_russian_tokens():
    fig, ax = plt.subplots(figsize=(9, 3))
    setup_style(ax)
    ax.set_xlim(-1, 10)
    ax.set_ylim(0, 1.2)
    ax.axis('off')

    tokens = [
        ("Щук", 0, 0.7), ("а", 0.8, 0.5), (" ", 1.4, 0.4),
        ("пой", 1.8, 0.7), ("мал", 2.5, 0.7), (" ", 3.1, 0.4),
        ("на", 3.5, 0.6), (" ", 4.0, 0.4),
        ("спин", 4.4, 0.7), ("нинг", 5.1, 0.7),
        (" ", 5.8, 0.4), ("в", 6.2, 0.5),
        (" ", 6.5, 0.4), ("Твер", 6.9, 0.7),
        ("ской", 7.6, 0.7), ("_", 8.3, 0.4),
        ("обл", 8.7, 0.7), ("асти", 9.2, 0.6),
    ]

    colors_pool = [COLORS['blue'], COLORS['purple'], COLORS['teal'],
                   COLORS['green'], COLORS['orange'], COLORS['accent']]

    for i, (tok, x, w) in enumerate(tokens):
        color = colors_pool[i % len(colors_pool)]
        rect = FancyBboxPatch((x, 0.3), w, 0.5,
                              boxstyle="round,pad=0.02",
                              facecolor=color, alpha=0.8,
                              edgecolor='white', linewidth=0.5)
        ax.add_patch(rect)
        ax.text(x + w/2, 0.55, tok, ha='center', va='center',
                fontsize=9, color='white', fontweight='bold')

    # Исходная строка сверху
    ax.text(5, 1.0, '"Щука поймал на спиннинг в Тверской области" → 18 токенов',
            ha='center', fontsize=9, color=COLORS['orange'], fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. 3: Пример BPE-токенизации рыбацкого текста',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 4: Tokenizer pipeline
# ============================================================
def diagram_pipeline():
    fig, ax = plt.subplots(figsize=(9, 3.5))
    setup_style(ax)
    ax.set_xlim(-1, 10)
    ax.set_ylim(0, 2.5)
    ax.axis('off')

    steps_data = [
        (0.0, 'Сбор\nкорпуса', 'Форумы, статьи,\nкниги о рыбалке', COLORS['green']),
        (2.5, 'Пре-процесс\nчистка', 'Удаление HTML,\nнормализация юникода', COLORS['blue']),
        (5.0, 'Обучение\nBPE', 'Слияние частых пар\nдо vocab_size', COLORS['purple']),
        (7.5, 'Токенизация\nтекста', 'Разбивка всех\nданных в IDs', COLORS['orange']),
    ]

    for x, title, desc, color in steps_data:
        rect = FancyBboxPatch((x, 1.1), 1.8, 0.9,
                              boxstyle="round,pad=0.04",
                              facecolor=color, alpha=0.2,
                              edgecolor=color, linewidth=2)
        ax.add_patch(rect)
        ax.text(x + 0.9, 1.55, title, ha='center', va='center',
                fontsize=9, color='white', fontweight='bold')
        ax.text(x + 0.9, 0.85, desc, ha='center', va='top',
                fontsize=7, color=COLORS['text'])

        if x > 0:
            ax.annotate('', xy=(x, 1.55), xytext=(x - 0.2, 1.55),
                        arrowprops=dict(arrowstyle='->', color=COLORS['accent'], lw=2.5))

    # Выход
    rect = FancyBboxPatch((4, 2.1), 1.5, 0.25,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['accent'], alpha=0.8,
                          edgecolor='white')
    ax.add_patch(rect)
    ax.text(4.75, 2.225, 'Tokenizer.model → файл', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. 4: Pipeline обучения и использования токенизатора',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 5: Vocabulary распределение
# ============================================================
def diagram_vocab_dist():
    fig, ax = plt.subplots(figsize=(7, 4))
    setup_style(ax)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)

    # Словарь — гистограмма
    bins = [10, 20, 30, 20, 10, 5, 3, 1, 0.5, 0.2]
    x_pos = np.arange(5, 100, 10)
    colors_bar = [COLORS['blue'], COLORS['green'], COLORS['orange'],
                  COLORS['accent'], COLORS['purple']] * 2
    for i, (h, c) in enumerate(zip(bins, colors_bar[:len(bins)])):
        ax.bar(x_pos[i], h, width=8, color=c, alpha=0.7, edgecolor='white', linewidth=0.5)

    ax.set_xlabel('Ранг токена (по частоте)', color=COLORS['text_dim'], fontsize=8)
    ax.set_ylabel('Частота использования (%)', color=COLORS['text_dim'], fontsize=8)
    ax.set_title('Распределение Zipf: немногие токены покрывают бóльшую часть текста',
                 color=COLORS['text'], fontsize=9)

    fig.text(0.5, 0.02, 'Рис. 5: Распределение частот токенов (закон Ципфа)',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 6: Special tokens
# ============================================================
def diagram_special_tokens():
    fig, ax = plt.subplots(figsize=(8, 3))
    setup_style(ax)
    ax.set_xlim(-1, 8)
    ax.set_ylim(0, 2.5)
    ax.axis('off')

    tokens_info = [
        ('<|endoftext|>', 'EOS\nконец текста', 'Модель останавливает\генерацию', COLORS['accent']),
        ('<|user|>', 'USER\nметка', 'Отделяет вопрос\nпользователя', COLORS['green']),
        ('<|assistant|>', 'ASSISTANT\nметка', 'Отделяет ответ\nмодели', COLORS['blue']),
        ('<|system|>', 'SYSTEM\nметка', 'Системный промпт\n(настройка)', COLORS['purple']),
    ]

    for i, (token, name, desc, color) in enumerate(tokens_info):
        x = i * 2.0
        rect = FancyBboxPatch((x, 1.0), 1.6, 0.8,
                              boxstyle="round,pad=0.03",
                              facecolor=color, alpha=0.2,
                              edgecolor=color, linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + 0.8, 1.6, token, ha='center', va='center',
                fontsize=8, color='white', fontweight='bold')
        ax.text(x + 0.8, 1.1, f'{name}\n{desc}', ha='center', va='center',
                fontsize=6, color=COLORS['text'])

    # Заголовок
    ax.text(4, 2.2, 'Special tokens — служебные токены в словаре',
            ha='center', fontsize=10, color=COLORS['orange'], fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. 6: Служебные токены в ChatML-формате',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ГЕНЕРАЦИЯ ДОКУМЕНТА
# ============================================================
def create_docx():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Титул
    title = doc.add_heading('Лекция 2: Tokenization — как LLM видит текст', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run(
        'FishingLLM — учим модель понимать рыбацкий язык\n'
        'Практика: обучение своего токенизатора на корпусе рыболовных текстов'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # === 1. Зачем нужна токенизация ===
    doc.add_heading('1. Зачем нужна токенизация', level=1)
    p = doc.add_paragraph()
    p.add_run('Текст — это строка символов. Нейросеть работает с числами.')
    p = doc.add_paragraph()
    p.add_run('Задача токенизации: ')
    run = p.add_run('превратить текст в последовательность целых чисел (ID), '
                     'с которыми может работать модель.')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('Почему это сложно:\n')
    items = [
        'Слов в языке много — словарь в 200K+ слов не поместится в VRAM\n'
        'Символов мало (~150) — но последовательность будет слишком длинной\n'
        'Нужен баланс: длина последовательности vs размер словаря'
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_picture(diagram_tokenizer_comparison(), width=Inches(5.5))
    doc.add_page_break()

    # === 2. BPE — Byte Pair Encoding ===
    doc.add_heading('2. BPE — Byte Pair Encoding', level=1)

    p = doc.add_paragraph()
    p.add_run('BPE (Byte Pair Encoding) — ')
    run = p.add_run('стандарт де-факто для современных LLM')
    run.bold = True
    p.add_run(' (GPT, LLaMA, Mistral, Qwen).')
    p = doc.add_paragraph()
    p.add_run('Идея: ').bold = True
    p.add_run('начинаем с отдельных символов, затем итеративно '
              'сливаем самую частую пару соседних токенов в один новый токен.')

    doc.add_picture(diagram_bpe_steps(), width=Inches(5.5))
    doc.add_page_break()

    doc.add_heading('2.1 Пример BPE на рыбацком корпусе', level=2)
    p = doc.add_paragraph()
    p.add_run('Допустим, корпус состоит из слов: ')
    p.add_run('"щука", "щуку", "щукой", "лещ", "леща", "окунь", "окуня"').italic = True

    code = doc.add_paragraph()
    run = code.add_run(
        '# Шаг 1: разбиваем на символы\n'
        'щ у к а _ щ у к у _ щ у к о й _ л е щ _ л е щ а _ о к у н ь _ о к у н я\n\n'
        '# Шаг 2: считаем частоту пар\n'
        '("щ","у"): 3, ("у","к"): 3, ("к","а"): 1, ("к","у"): 1, ("к","о"): 1, ...\n\n'
        '# Шаг 3: сливаем самую частую пару "щ"+"у" → "щу"\n'
        'щу к а _ щу к у _ щу к о й _ л е щ _ л е щ а _ о к у н ь _ о к у н я\n\n'
        '# Шаг 4: повторяем... в итоге словарь будет содержать:\n'
        '["щук", "а", "у", "ой", "лещ", "окун", "ь", "я", "_"]'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run('Слово "щука" будет представлено как [щук, а] — 2 токена вместо 4 символов.')
    run.font.italic = True

    doc.add_picture(diagram_russian_tokens(), width=Inches(6))
    doc.add_page_break()

    # === 3. SentencePiece (Unigram) ===
    doc.add_heading('3. SentencePiece — альтернатива BPE', level=1)

    p = doc.add_paragraph()
    p.add_run('SentencePiece (Google, 2018) — ')
    run = p.add_run('используется в LLaMA, Qwen')
    run.bold = True
    p.add_run('. Отличие от BPE: работает напрямую с сырым текстом (без предварительного деления на слова по пробелам).')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Shading Accent 1'
    cells = [
        ('Характеристика', 'BPE (GPT, Mistral)', 'SentencePiece (LLaMA, Qwen)'),
        ('Единица', 'Байты → токены', 'Unicode символы → токены'),
        ('Пробелы', 'Пробел = отдельный токен', 'Пробел = "_" (часть токена)'),
        ('Алгоритм', 'Слияние частых пар', 'Unigram LM (вероятностный)'),
    ]
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            table.cell(i, j).text = val
            if i == 0:
                for paragraph in table.cell(i, j).paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_page_break()

    # === 4. Практика ===
    doc.add_heading('4. Практика: токенизируем наш корпус', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Будем использовать библиотеку tokenizers (Hugging Face). '
                     'Вот как выглядит обучение токенизатора на наших данных:')
    run.bold = True

    code = doc.add_paragraph()
    run = code.add_run(
        'from tokenizers import Tokenizer\n'
        'from tokenizers.models import BPE\n'
        'from tokenizers.trainers import BpeTrainer\n'
        'from tokenizers.pre_tokenizers import ByteLevel\n\n'
        '# Создаём токенизатор\n'
        'tokenizer = Tokenizer(BPE(unk_token="<unk>"))\n'
        'tokenizer.pre_tokenizer = ByteLevel()\n\n'
        '# Обучаем на рыбацком корпусе\n'
        'trainer = BpeTrainer(\n'
        '    vocab_size=16_000,        # размер словаря\n'
        '    special_tokens=["<pad>", "<unk>", "<s>", "</s>",\n'
        '                   "<|user|>", "<|assistant|>"],\n'
        '    min_frequency=2,\n'
        ')\n\n'
        'files = ["fishing_corpus.txt"]\n'
        'tokenizer.train(files, trainer)\n\n'
        '# Пробуем\n'
        'output = tokenizer.encode("щука на спиннинг")\n'
        'print(output.tokens)   # ["щу", "ка", "_", "на", "_", "спин", "нинг"]\n'
        'print(output.ids)      # [341, 890, 307, 56, 307, 4512, 3781]'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

    p = doc.add_paragraph()
    p.add_run('Важные параметры:\n').bold = True

    params = [
        'vocab_size — чем больше, тем лучше покрытие, но тяжелее embedding layer\n'
        'special_tokens — служебные токены для формата диалога (ChatML)\n'
        'min_frequency — минимальная частота пары для слияния (фильтрует шум)'
    ]
    for param in params:
        p = doc.add_paragraph(param, style='List Bullet')

    doc.add_picture(diagram_pipeline(), width=Inches(5.5))
    doc.add_page_break()

    # === 5. Special Tokens ===
    doc.add_heading('5. Special Tokens — служебные токены', level=1)

    p = doc.add_paragraph()
    p.add_run('Когда мы будем делать fine-tuning или обучение с нуля, '
              'формат диалога будет выглядеть так:')

    code = doc.add_paragraph()
    run = code.add_run(
        '<|system|>\n'
        'Ты эксперт по рыбалке в Тверской области.\n'
        '<|user|>\n'
        'Какая наживка лучше для леща в июне?\n'
        '<|assistant|>\n'
        'В июне на Верхней Волге лещ хорошо берёт на ...'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('Каждый <|...|> — это ОДИН токен в словаре, а не 4 символа. '
              'Модель учится понимать их как служебные метки, а не часть текста.')

    doc.add_picture(diagram_special_tokens(), width=Inches(5.5))
    doc.add_page_break()

    # === 6. Закон Ципфа ===
    doc.add_heading('6. Закон Ципфа и распределение токенов', level=1)

    p = doc.add_paragraph()
    p.add_run('В любом естественном языке немногие слова покрывают бóльшую часть текста. '
              'Те же 10% самых частых слов покрывают ~80% всех употреблений.')
    p = doc.add_paragraph()
    p.add_run('Для токенов то же самое: ')
    p.add_run('первые несколько тысяч токенов покрывают бóльшую часть текста, '
              'а остальные 10K+ — редкие слова и подслова.')
    p.add_run(' Это называется закон Ципфа (Zipf law).')

    doc.add_picture(diagram_vocab_dist(), width=Inches(5))
    doc.add_page_break()

    # === 7. Итог ===
    doc.add_heading('7. Итог: что нужно запомнить', level=1)

    items = [
        'Токенизация — первый этап работы LLM: текст → числа',
        'BPE — итеративное слияние частых пар символов/токенов',
        'SentencePiece — вариант BPE, работающий с Unicode, без предварительного деления на слова',
        'Размер словаря (vocab_size) — гиперпараметр: 16K для маленьких моделей, 128K для больших',
        'Special tokens — служебные токены для разметки диалога',
        'Токенизатор НЕ меняется после обучения модели',
        'Для русского языка важно выбирать токенизатор, обученный на кириллице (или дообучать свой)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Number')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Следующий шаг: Лекция 3 — Embeddings и представление текста')
    run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run('Практическое задание: обучить свой первый BPE-токенизатор на русском корпусе')
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, 'lecture-02-tokenization.docx')
    doc.save(path)
    print(f'Документ сохранён: {path}')


if __name__ == '__main__':
    create_docx()
