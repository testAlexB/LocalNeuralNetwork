"""Лекция 1: Transformer — полное, университетское изложение
Основа: Vaswani et al. (2017), Hochreiter & Schmidhuber (1997), Karpathy (2023)
Цель: дать студенту 3-го курса РАЗОБРАТЬСЯ, а не "пробежать глазами"
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io, os, sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')

# Dark theme color palette
C = {
    'bg': '#1a1a2e', 'primary': '#16213e', 'secondary': '#0f3460',
    'accent': '#e94560', 'orange': '#f59e0b', 'green': '#10b981',
    'blue': '#3b82f6', 'purple': '#8b5cf6', 'teal': '#06b6d4',
    'pink': '#ec4899', 'text': '#e2e8f0', 'text_dim': '#94a3b8',
}

def style_ax(ax):
    ax.set_facecolor(C['bg'])
    ax.figure.patch.set_facecolor(C['bg'])
    for s in ax.spines.values():
        s.set_visible(False)
    ax.tick_params(colors=C['text_dim'], labelsize=6)
    ax.grid(True, alpha=0.06, color=C['text_dim'])

def save_fig(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=160, bbox_inches='tight',
                facecolor=C['bg'], edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 1: RNN — sequential processing problem
# ═══════════════════════════════════════════════════════════════
def fig_rnn_vs_transformer():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 3.5))
    for ax in [ax1, ax2]:
        style_ax(ax)
        ax.axis('off')

    # RNN
    ax1.text(0, 0.85, 'RNN / LSTM (последовательная обработка)',
             fontsize=9, color=C['accent'], fontweight='bold', ha='center')
    tokens = ['Щ', 'у', 'к', 'а', ' ', 'л', 'е', 'щ']
    for i, tok in enumerate(tokens):
        x = i * 0.85
        rect = FancyBboxPatch((x, 0.3), 0.7, 0.35,
                              boxstyle="round,pad=0.03",
                              facecolor=C['primary'], alpha=0.8,
                              edgecolor=C['blue'], linewidth=1)
        ax1.add_patch(rect)
        ax1.text(x + 0.35, 0.475, tok, ha='center', va='center',
                 fontsize=10, color='white', fontweight='bold')
        if i > 0:
            ax1.annotate('', xy=(x, 0.475), xytext=(x - 0.15, 0.475),
                         arrowprops=dict(arrowstyle='->', color=C['accent'], lw=2))
    ax1.text(3.0, 0.1, 'Шаг за шагом: токен i ждёт hidden_state от i-1\n→ O(n) времени, забывает начало',
             fontsize=7, color=C['text_dim'], ha='center')

    # Transformer
    ax2.text(0.5, 0.85, 'Transformer (параллельная обработка)',
             fontsize=9, color=C['green'], fontweight='bold', ha='center')
    for i, tok in enumerate(tokens):
        x = i * 0.85 + 0.5
        rect = FancyBboxPatch((x, 0.3), 0.7, 0.35,
                              boxstyle="round,pad=0.03",
                              facecolor=C['primary'], alpha=0.8,
                              edgecolor=C['green'], linewidth=1)
        ax2.add_patch(rect)
        ax2.text(x + 0.35, 0.475, tok, ha='center', va='center',
                 fontsize=10, color='white', fontweight='bold')

    # Attention connections from "л" and "е" back to "щ"
    # "лещ" word at positions
    for src in [0, 4, 5, 6, 7]:  # various connections
        for dst in [4, 5, 6, 7]:
            if src != dst:
                x1 = src * 0.85 + 0.85
                y1 = 0.475
                x2 = dst * 0.85 + 0.85
                y2 = 0.475
                dx = x2 - x1
                dy = y2 - y1
                ax2.annotate('', xy=(x2, y2), xytext=(x1, y1),
                             arrowprops=dict(arrowstyle='->', color=C['purple'],
                                             lw=0.5, alpha=0.3))
    ax2.text(4.0, 0.1, 'Все токены обрабатываются одновременно\n→ O(1) шагов, каждый видит всех\n→ Self-Attention: "посмотри на всё"',
             fontsize=7, color=C['text_dim'], ha='center')

    fig.text(0.5, 0.02, 'Рис. 1.2: Последовательная (RNN) vs параллельная (Transformer) обработка',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 2: RNN cell, vanishing gradient
# ═══════════════════════════════════════════════════════════════
def fig_rnn_cell():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
    for ax in [ax1, ax2]:
        style_ax(ax)
        ax.axis('off')

    # Left: RNN cell
    for i, tok in enumerate(['x₁', 'x₂', 'x₃', '...', 'xₙ']):
        x = i * 1.8
        # Input circle
        circ = plt.Circle((x, 0.6), 0.3, color=C['blue'], alpha=0.7)
        ax1.add_patch(circ)
        ax1.text(x, 0.6, tok, ha='center', va='center', fontsize=8, color='white')
        # RNN cell
        cell = FancyBboxPatch((x - 0.5, 0.0), 1.0, 0.45,
                              boxstyle="round,pad=0.05",
                              facecolor=C['secondary'], alpha=0.9,
                              edgecolor=C['orange'], linewidth=1.5)
        ax1.add_patch(cell)
        ax1.text(x, 0.225, f'h_{i+1} = tanh(W·[h_{i}, x_{i+1}])',
                 ha='center', va='center', fontsize=6, color=C['text'])
        # Arrow right
        if i < 4:
            ax1.annotate('', xy=(x + 0.8, 0.225), xytext=(x + 0.5, 0.225),
                         arrowprops=dict(arrowstyle='->', color=C['accent'], lw=2))

    ax1.text(4.5, -0.25, 'hₙ зависит от h₁ через n умножений матриц\n'
                         'gradient ∝ Wⁿ → при n>10 градиент взрывается или затухает',
             fontsize=7, color=C['text_dim'], ha='center')

    # Right: LSTM cell
    components = [
        (0.0, C['green'], 'Forget Gate\nσ(W·[h_{t-1}, x_t])'),
        (1.5, C['blue'], 'Input Gate\nσ(W·[h_{t-1}, x_t])'),
        (3.0, C['purple'], 'Cell State\nC_t = f_t·C_{t-1} + i_t·tanh(...)'),
        (4.5, C['orange'], 'Output Gate\nσ(W·[h_{t-1}, x_t])'),
    ]
    for x, color, label in components:
        rect = FancyBboxPatch((x, 0.3), 1.2, 0.5,
                              boxstyle="round,pad=0.03",
                              facecolor=color, alpha=0.2,
                              edgecolor=color, linewidth=1.5)
        ax2.add_patch(rect)
        ax2.text(x + 0.6, 0.55, label, ha='center', va='center',
                 fontsize=6, color=C['text'])

    # Cell state line through gates
    ax2.annotate('', xy=(0, 0.55), xytext=(5.5, 0.55),
                 arrowprops=dict(arrowstyle='->', color=C['accent'], lw=3))
    ax2.text(5.8, 0.55, 'C_t (cell state)', fontsize=7, color=C['accent'],
             fontweight='bold', va='center')

    ax2.text(2.8, 0.0, 'LSTM решает проблему затухания градиента через "конвейер" C_t\n'
                       'НО: всё ещё последовательный, не параллелится',
             fontsize=7, color=C['text_dim'], ha='center')

    fig.text(0.5, 0.02, 'Рис. 1.1: RNN и LSTM — принцип работы и проблема последовательности',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 3: Attention mechanism — QKV visual
# ═══════════════════════════════════════════════════════════════
def fig_attention_qkv():
    fig, ax = plt.subplots(figsize=(9, 5.5))
    style_ax(ax)
    ax.set_xlim(-1, 5)
    ax.set_ylim(0, 3.5)
    ax.axis('off')

    # Title
    ax.text(2.5, 3.2, 'Внимание (Attention) = запрос в базу данных',
            ha='center', fontsize=11, color=C['accent'], fontweight='bold')

    # Three columns: Q, K, V
    cols = [
        ('Query (Q) — "что я ищу?"', 0.0, C['blue']),
        ('Key (K) — "что у меня есть?"', 1.8, C['purple']),
        ('Value (V) — "информация"', 3.6, C['green']),
    ]
    for title, x, color in cols:
        ax.text(x + 0.5, 2.6, title, fontsize=8, color=color,
                fontweight='bold', ha='center')
        # Draw boxes representing vectors
        for row in range(5):
            for col in range(3):
                val = np.random.randint(0, 10)
                alpha = 0.25 + 0.65 * (val / 10)
                r = FancyBboxPatch((x + col * 0.25, 2.0 - row * 0.25),
                                   0.22, 0.22, boxstyle="round,pad=0.01",
                                   facecolor=color, alpha=alpha,
                                   edgecolor='white', linewidth=0.3)
                ax.add_patch(r)
                ax.text(x + col * 0.25 + 0.11, 2.0 - row * 0.25 + 0.11,
                        str(val), ha='center', va='center', fontsize=5, color='white')

    # Down arrow and formula
    ax.text(2.5, 1.1, 'Аналогия: поиск в Google\n'
                      'Вы вводите запрос (Query)\n'
                      'Google ищет по ключевым словам (Keys)\n'
                      'И возвращает содержимое (Values)',
            ha='center', fontsize=8, color=C['text'], va='center')

    # The big formula
    formula_box = FancyBboxPatch((0.2, 0.05), 4.6, 0.55,
                                 boxstyle="round,pad=0.05",
                                 facecolor=C['secondary'], alpha=0.9,
                                 edgecolor=C['orange'], linewidth=1.5)
    ax.add_patch(formula_box)
    ax.text(2.5, 0.325, 'Attention(Q, K, V) = softmax( Q · Kᵀ / √dₖ ) · V',
            ha='center', va='center', fontsize=11, color=C['orange'],
            fontweight='bold', family='monospace')

    # Legend
    ax.text(0.05, 0.78, 'Q — матрица [seq_len × dₖ], где каждый токен "спрашивает"\n'
                        'K — матрица [seq_len × dₖ], где каждый токен "отвечает"\n'
                        'V — матрица [seq_len × dᵥ], где каждый токен "передаёт"\n'
                        '√dₖ — масштабирование (чтобы не взрывались softmax)\n'
                        'softmax — превращает веса в вероятности',
            fontsize=6.5, color=C['text'])

    fig.text(0.5, 0.01, 'Рис. 2.1: Self-Attention — механизм QKV (Vaswani et al., 2017)',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 4: Attention weight matrix visualization
# ═══════════════════════════════════════════════════════════════
def fig_attention_weights():
    fig, ax = plt.subplots(figsize=(6, 5))
    style_ax(ax)

    words = ['Щука', 'поймана', 'на', 'спиннинг', 'в', 'июне']
    n = len(words)

    # Synthesize attention weights
    rng = np.random.RandomState(42)
    W = rng.rand(n, n)
    # Make diagonal and fishing-related pairs stronger
    for i in range(n):
        W[i, i] *= 2
    W[0, 1] *= 1.8  # щука-поймана
    W[1, 0] *= 1.8
    W[3, 4] *= 1.5  # спиннинг-в
    W[4, 5] *= 1.5  # в-июне
    W = W / W.sum(axis=1, keepdims=True)

    im = ax.imshow(W, cmap='viridis', aspect='auto', vmin=0, vmax=W.max())

    ax.set_xticks(range(n))
    ax.set_yticks(range(n))
    ax.set_xticklabels(words, fontsize=9, color=C['text'])
    ax.set_yticklabels(words, fontsize=9, color=C['text'])
    ax.tick_params(colors=C['text_dim'], labelsize=6)

    # Add text annotations
    for i in range(n):
        for j in range(n):
            color = 'white' if W[i, j] > 0.25 else C['text']
            ax.text(j, i, f'{W[i, j]:.2f}', ha='center', va='center',
                    fontsize=7, color=color)

    ax.set_xlabel('Key (на что смотрят)', fontsize=8, color=C['text'])
    ax.set_ylabel('Query (кто смотрит)', fontsize=8, color=C['text'])
    ax.set_title('Матрица внимания: ярче = сильнее внимание',
                 fontsize=9, color=C['text'], fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. 2.2: Матрица попарного внимания для фразы "Щука поймана на спиннинг в июне"',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 5: Multi-Head Attention
# ═══════════════════════════════════════════════════════════════
def fig_multihead():
    fig, ax = plt.subplots(figsize=(8, 5))
    style_ax(ax)
    ax.set_xlim(-1, 8)
    ax.set_ylim(0, 4)
    ax.axis('off')

    # Input
    inp = FancyBboxPatch((2.8, 3.3), 2.0, 0.4,
                         boxstyle="round,pad=0.04",
                         facecolor=C['green'], alpha=0.85,
                         edgecolor='white')
    ax.add_patch(inp)
    ax.text(3.8, 3.5, 'INPUT (X)', ha='center', va='center',
            fontsize=10, color='white', fontweight='bold')

    # Split to heads
    heads = 5
    for i in range(heads):
        x = i * 1.2
        # Head circle
        circ = plt.Circle((x + 0.6, 2.3), 0.5, color=C['purple'],
                          alpha=0.3 + i * 0.12, ec=C['purple'], lw=1.5)
        ax.add_patch(circ)
        ax.text(x + 0.6, 2.3, f'Head\n{i+1}', ha='center', va='center',
                fontsize=7, color='white', fontweight='bold')

        # Arrows from input to each head
        ax.annotate('', xy=(x + 0.6, 2.8), xytext=(3.8, 3.3),
                    arrowprops=dict(arrowstyle='->', color=C['text'], lw=1, alpha=0.4))
        # Arrows down to concat
        ax.annotate('', xy=(x + 0.6, 1.1), xytext=(x + 0.6, 1.8),
                    arrowprops=dict(arrowstyle='->', color=C['text'], lw=1, alpha=0.4))

    # Concat
    concat = FancyBboxPatch((1.5, 0.6), 4.0, 0.35,
                            boxstyle="round,pad=0.04",
                            facecolor=C['orange'], alpha=0.85,
                            edgecolor='white')
    ax.add_patch(concat)
    ax.text(3.5, 0.775, 'CONCATENATE (склейка)', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Output projection
    out = FancyBboxPatch((2.3, 0.1), 2.4, 0.35,
                         boxstyle="round,pad=0.04",
                         facecolor=C['accent'], alpha=0.85,
                         edgecolor='white')
    ax.add_patch(out)
    ax.text(3.5, 0.275, 'Linear Projection', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Right side text
    ax.text(7.0, 2.3, 'Каждая голова\nищет свой\nтип связей:\n\n'
                      '• Head 1: синтаксис\n'
                      '• Head 2: семантика\n'
                      '• Head 3: числа\n'
                      '• Head 4: локации\n'
                      '• Head 5: время',
            fontsize=7, color=C['text'])

    # Formula
    ax.text(3.5, 3.9, 'MultiHead(Q,K,V) = Concat(head₁ ... headₕ) · Wᵒ',
            ha='center', fontsize=8, color=C['orange'], fontweight='bold',
            family='monospace')

    fig.text(0.5, 0.02, 'Рис. 3: Multi-Head Attention — параллельное внимание (Vaswani et al., 2017)',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 6: Positional Encoding
# ═══════════════════════════════════════════════════════════════
def fig_positional():
    fig, ax = plt.subplots(figsize=(8, 4))
    style_ax(ax)

    # Sin curves for different positions
    pos = np.arange(50)
    freqs = [0.01, 0.05, 0.1, 0.2, 0.5]
    colors_pos = [C['blue'], C['green'], C['orange'], C['accent'], C['purple']]
    d_labels = ['d=0', 'd=10', 'd=20', 'd=30', 'd=40']

    for freq, color, label in zip(freqs, colors_pos, d_labels):
        vals = np.sin(pos * freq)
        ax.plot(vals, color=color, alpha=0.8, linewidth=1.5, label=label)

    ax.set_xlabel('Позиция слова в предложении', fontsize=8, color=C['text'])
    ax.set_ylabel('Значение позиционного сигнала', fontsize=8, color=C['text'])
    ax.set_title('Sinusoidal Positional Encoding:\nкаждая размерность имеет свою "частоту"',
                 fontsize=9, color=C['text'], fontweight='bold')
    ax.legend(fontsize=7, loc='upper right',
              facecolor=C['primary'], edgecolor=C['text_dim'],
              labelcolor=C['text'])
    ax.spines['left'].set_visible(True)
    ax.spines['bottom'].set_visible(True)

    # Formula on the side
    fig.text(0.85, 0.5, 'PE(pos, 2i) = sin(pos / 10000^{2i/d})\n'
                        'PE(pos, 2i+1) = cos(pos / 10000^{2i/d})',
             ha='center', fontsize=8, color=C['orange'],
             fontweight='bold', family='monospace')

    fig.text(0.5, 0.02, 'Рис. 4: Синусоидальное позиционное кодирование (Vaswani et al., 2017)',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 7: Transformer Decoder Layer
# ═══════════════════════════════════════════════════════════════
def fig_decoder_layer():
    fig, ax = plt.subplots(figsize=(5.5, 7))
    style_ax(ax)
    ax.set_xlim(-1, 1)
    ax.set_ylim(-1, 1.5)
    ax.axis('off')

    layers = [
        (-0.35, 1.1, 0.7, 0.1, 'ВХОД (X + PE)', C['green'], True),
        (-0.30, 0.85, 0.6, 0.1, 'Layer Norm', C['teal'], True),
        (-0.35, 0.55, 0.7, 0.18, 'Multi-Head\nSelf-Attention', C['purple'], True),
        (-0.30, 0.25, 0.6, 0.1, 'Layer Norm', C['teal'], True),
        (-0.35, -0.05, 0.7, 0.18, 'Feed-Forward\nNetwork (FFN)', C['blue'], True),
        (-0.25, -0.35, 0.5, 0.1, 'ВЫХОД (→ слой N+1)', C['accent'], True),
    ]

    for x, y, w, h, label, color, arrow in layers:
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.03",
                              facecolor=color, alpha=0.85,
                              edgecolor='white', linewidth=0.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label, ha='center', va='center',
                fontsize=7, color='white', fontweight='bold')

    # Arrows down
    for i, (_, y, _, h, _, _, arrow) in enumerate(layers[:-1]):
        _, y_next, _, _, _, _, _ = layers[i + 1]
        ax.annotate('', xy=(0, y_next + y_next), xytext=(0, y),
                    arrowprops=dict(arrowstyle='->', color=C['text'], lw=1.5))

    # Residual connections (dashed arrows going right then down)
    # From input to after attention
    ax.annotate('', xy=(0.38, 0.55), xytext=(0.38, 1.1),
                arrowprops=dict(arrowstyle='->', color=C['orange'],
                                lw=1.5, linestyle='dashed'))
    ax.text(0.42, 0.82, '+', fontsize=10, color=C['orange'], fontweight='bold')

    # From after attention to after FFN
    ax.annotate('', xy=(0.38, -0.05), xytext=(0.38, 0.55),
                arrowprops=dict(arrowstyle='->', color=C['orange'],
                                lw=1.5, linestyle='dashed'))
    ax.text(0.42, 0.25, '+', fontsize=10, color=C['orange'], fontweight='bold')

    # Legend right side
    ax.text(0.65, 0.85, 'x_out = x_in + Layer(x_in)\n'
                        'Residual connection:\n'
                        'если слой не дал\n'
                        'прироста — данные\n'
                        'проходят насквозь',
            fontsize=6, color=C['orange'])

    fig.text(0.5, 0.02, 'Рис. 5: Один слой Decoder-Transformer с residual-связями',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# DIAGRAM 8: Generation process with probability
# ═══════════════════════════════════════════════════════════════
def fig_generation():
    fig, ax = plt.subplots(figsize=(9, 4))
    style_ax(ax)
    ax.set_xlim(-0.5, 10)
    ax.set_ylim(0, 1.2)
    ax.axis('off')

    # Input box
    inp = FancyBboxPatch((0.0, 0.55), 3.0, 0.35,
                         boxstyle="round,pad=0.03",
                         facecolor=C['green'], alpha=0.85,
                         edgecolor='white')
    ax.add_patch(inp)
    ax.text(1.5, 0.725, 'ВХОД: "Какая наживка\nлучше для леща в..."',
            ha='center', va='center', fontsize=8, color='white', fontweight='bold')

    # Transformer block
    tr = FancyBboxPatch((3.5, 0.5), 1.8, 0.45,
                        boxstyle="round,pad=0.03",
                        facecolor=C['secondary'], alpha=0.9,
                        edgecolor=C['accent'], linewidth=2)
    ax.add_patch(tr)
    ax.text(4.4, 0.725, 'Transformer\nDecoder', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Arrow to prob distribution
    ax.annotate('', xy=(5.5, 0.725), xytext=(5.3, 0.725),
                arrowprops=dict(arrowstyle='->', color=C['text'], lw=2))

    # Probability distribution
    probs = {'июне': 0.31, 'августе': 0.18, 'июле': 0.22,
             'мае': 0.12, 'сентябре': 0.08, 'апреле': 0.05,
             'октябре': 0.03, 'ноябре': 0.01}

    x_start = 5.8
    for i, (tok, prob) in enumerate(probs.items()):
        x = x_start + i * 0.42
        h = prob * 4
        is_best = tok == 'июне'
        color = C['orange'] if is_best else C['blue']
        alpha = 0.8 if is_best else 0.3 + prob * 2
        rect = FancyBboxPatch((x - 0.12, 0.08), 0.24, h,
                              boxstyle="round,pad=0.01",
                              facecolor=color, alpha=alpha,
                              edgecolor='white', linewidth=0.5 if not is_best else 1.5)
        ax.add_patch(rect)
        ax.text(x, h + 0.1, f'{prob:.0%}', ha='center', va='bottom',
                fontsize=6, color=C['text'])
        ax.text(x, 0.02, tok, ha='center', va='bottom',
                fontsize=6, color=C['text'])

    # Highlight chosen token
    x_hl = x_start + 0
    rect = FancyBboxPatch((x_hl - 0.18, 0.44), 0.36, 0.35,
                          boxstyle="round,pad=0.02",
                          facecolor=C['orange'], alpha=0.2,
                          edgecolor=C['orange'], linewidth=2, linestyle='dashed')
    ax.add_patch(rect)
    ax.text(x_hl, 0.85, 'Выбран\n"июне"\n(наибольшая\nвероятность)',
            ha='center', va='center', fontsize=5.5, color=C['orange'],
            fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. 6: Авторегрессивная генерация — предсказание следующего токена',
             ha='center', fontsize=8, color=C['text_dim'])
    return save_fig(fig)


# ═══════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════════════════════════
def build_doc():
    doc = Document()

    # Global style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ═══ TITLE ═══
    title = doc.add_heading('Лекция 1: Transformer — архитектура, изменившая NLP', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'FishingLLM — Создание экспертной LLM по рыбалке в Тверской области\n'
        'Основано на: Vaswani et al. (2017), Hochreiter & Schmidhuber (1997),\n'
        'Bahdanau et al. (2015), Karpathy (2023), 3Blue1Brown (2024)'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # ═══ INTRODUCTION ═══
    doc.add_heading('Введение: почему мы изучаем Transformer?', level=1)

    p = doc.add_paragraph()
    p.add_run('Наша цель — создать нейросетевую модель, которая будет экспертом по рыбалке '
              'в Тверской области. Для этого нам нужно понимать, как работает "мозг" современной LLM. '
              'А работает она на архитектуре Transformer.')

    p = doc.add_paragraph()
    p.add_run('Прежде чем погрузиться в детали, давайте ответим на главный вопрос: ')
    run = p.add_run('почему Transformer, а не что-то другое?')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

    p = doc.add_paragraph()
    p.add_run('Чтобы ответить, нужно понять, что было ДО. И почему это "до" перестало устраивать.')
    doc.add_page_break()

    # ═══ 1. HISTORICAL CONTEXT ═══
    doc.add_heading('Глава 1. До Transformer: эпоха рекуррентных нейросетей', level=1)

    doc.add_heading('1.1. RNN — первая попытка обрабатывать последовательности', level=2)

    p = doc.add_paragraph()
    p.add_run('Рекуррентные нейронные сети (RNN) были предложены ещё в 1980-х годах (Rumelhart et al., 1986). '
              'Их ключевая идея: у сети есть "память" — скрытое состояние (hidden state), '
              'которое передаётся от одного шага к другому.')

    p = doc.add_paragraph()
    p.add_run('Формально: ')
    run = p.add_run('hₜ = tanh(W·[hₜ₋₁; xₜ] + b)')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p.add_run(', где hₜ — скрытое состояние на шаге t, xₜ — вход на шаге t.')

    p = doc.add_paragraph()
    p.add_run('На первый взгляд, всё логично: читаем слово за словом, сохраняя контекст. '
              'Но на практике RNN страдает от фундаментальной проблемы.')

    doc.add_heading('1.2. Проблема затухающих градиентов (Vanishing Gradient)', level=2)

    p = doc.add_paragraph()
    p.add_run('Рассмотрим предложение: "Щука, которую я поймал вчера на спиннинг возле устья Медведицы, '
              'весила 12 кг". Чтобы предсказать слово "весила", модель должна "вспомнить" слово "Щука" '
              'из начала предложения. Между ними — 10+ слов.')

    p = doc.add_paragraph()
    p.add_run('При обучении RNN градиент (сигнал ошибки) распространяется обратно во времени. '
              'На каждом шаге он умножается на матрицу весов W. Если собственные числа W < 1, '
              'градиент экспоненциально затухает:')

    p = doc.add_paragraph()
    run = p.add_run('∂L/∂h₁ = ∂L/∂hₙ · Wⁿ')
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    p.add_run(', где n — расстояние между словами.')

    p = doc.add_paragraph()
    p.add_run('При n=10 и λ<1 градиент практически равен нулю. Модель не может учить '
              'долгосрочные зависимости. Это — ')
    run = p.add_run('проблема затухающих градиентов')
    run.bold = True
    p.add_run(' (Hochreiter, 1991; Bengio et al., 1994).')

    doc.add_picture(fig_rnn_cell(), width=Inches(5.5))
    doc.add_page_break()

    doc.add_heading('1.3. LSTM — попытка решить проблему', level=2)

    p = doc.add_paragraph()
    p.add_run('Hochreiter & Schmidhuber (1997) предложили Long Short-Term Memory (LSTM) — '
              'архитектуру с "конвейером" клеточного состояния Cₜ, который может переносить '
              'информацию через десятки шагов без затухания.')

    p = doc.add_paragraph()
    p.add_run('Ключевые формулы LSTM:')
    items = [
        'Забывающий шлюз: fₜ = σ(Wf·[hₜ₋₁, xₜ] + bf) — решает, что забыть из старой памяти',
        'Входной шлюз: iₜ = σ(Wi·[hₜ₋₁, xₜ] + bi) — решает, что записать в новую память',
        'Клеточное состояние: Cₜ = fₜ·Cₜ₋₁ + iₜ·tanh(Wc·[hₜ₋₁, xₜ] + bc) — обновление памяти',
        'Выходной шлюз: oₜ = σ(Wo·[hₜ₋₁, xₜ] + bo) — решает, что выдать на выход',
        'Скрытое состояние: hₜ = oₜ·tanh(Cₜ)',
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.runs[0].font.size = Pt(9)
        if len(p.runs) > 0:
            p.runs[0].font.name = 'Consolas'

    p = doc.add_paragraph()
    p.add_run('LSTM решила проблему затухания градиента, но ')
    run = p.add_run('не решила проблему последовательной обработки')
    run.bold = True
    p.add_run(': hₜ по-прежнему зависит от hₜ₋₁, а значит, мы не можем обрабатывать слова параллельно. '
              'Для обучения на больших корпусах это катастрофически медленно.')

    p = doc.add_paragraph()
    p.add_run('Важный вывод: ')
    run = p.add_run('к 2017 году сообщество осознало, что рекуррентный подход — тупиковый '
                     'для масштабирования. Нужна была архитектура, которая:')
    run.bold = True
    p.add_run(' 1) обрабатывает все слова параллельно; 2) умеет "смотреть" на любое слово '
              'в предложении без потери информации; 3) масштабируется на миллиарды параметров.')
    doc.add_page_break()

    # ═══ 2. ATTENTION ═══
    doc.add_heading('Глава 2. Внимание (Attention) — ключевая идея', level=1)

    doc.add_heading('2.1. Откуда пришла идея внимания', level=2)

    p = doc.add_paragraph()
    p.add_run('В 2015 году Bahdanau et al. предложили механизм внимания для машинного перевода. '
              'Идея была проста: вместо того чтобы сжимать всё предложение в один вектор (как делали '
              'Seq2Seq модели), давайте на каждом шаге перевода "заглядывать" в исходное предложение '
              'и выбирать, какие слова там важны.')

    p = doc.add_paragraph()
    p.add_run('Это называется ')
    run = p.add_run('additive attention')
    run.bold = True
    p.add_run(' (или Bahdanau attention). Но в 2017 году Vaswani et al. предложили более '
              'эффективный вариант — dot-product attention (скалярное произведение), '
              'который и стал стандартом.')

    doc.add_heading('2.2. Query, Key, Value — аналогия с поиском', level=2)

    p = doc.add_paragraph()
    p.add_run('Чтобы понять Q, K, V, представьте, что вы ищете информацию о ловле щуки:')

    p = doc.add_paragraph()
    p.add_run('• ')
    run = p.add_run('Query (Q) — это ваш поисковый запрос')
    run.bold = True
    p.add_run(': "ловля щуки на спиннинг"' + '\n')
    p.add_run('• ')
    run = p.add_run('Key (K) — это заголовки статей в базе знаний')
    run.bold = True
    p.add_run(': "Как ловить щуку", "Снасти для спиннинга", "Рыбалка в Тверской области"' + '\n')
    p.add_run('• ')
    run = p.add_run('Value (V) — это содержимое этих статей')
    run.bold = True

    p = doc.add_paragraph()
    p.add_run('Вы (Q) сравниваете свой запрос со всеми заголовками (K), выбираете наиболее '
              'релевантные (через softmax) и читаете их содержание (V). В нейросети то же самое — '
              'но на уровне токенов.')

    doc.add_heading('2.3. Формула внимания — разбор каждого символа', level=2)

    p = doc.add_paragraph()
    run = p.add_run('Attention(Q, K, V) = softmax( Q · Kᵀ / √dₖ ) · V')
    run.font.size = Pt(12)
    run.font.name = 'Consolas'
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('А теперь разберём, что здесь происходит, символ за символом:')

    items = [
        ('Q', 'Матрица Query размером [seq_len × dₖ]. Каждая строка — это "запрос" одного токена ко всем остальным.'),
        ('Kᵀ', 'Транспонированная матрица Key [dₖ × seq_len]. Каждый столбец — "ответ" токена на запросы.'),
        ('Q·Kᵀ', 'Матрица внимания [seq_len × seq_len]. Элемент (i,j) — скалярное произведение query токена i и key токена j. Чем больше — тем сильнее токен i "обращает внимание" на токен j.'),
        ('√dₖ', 'Масштабирование. dₖ — размерность ключа (обычно 64-128 для одной головы). Без деления на √dₖ значения скалярных произведений при большой размерности становятся слишком большими, и softmax превращается в "argmax" — теряется градиент.'),
        ('softmax', 'Функция, превращающая веса в вероятности (сумма по строкам = 1). После softmax токен "распределяет" своё внимание между всеми токенами.'),
        ('V', 'Матрица Value [seq_len × dᵥ]. То, что будет "возвращено" — информация, которую токены передают друг другу.'),
        ('· V', 'Умножение матрицы внимания [seq_len × seq_len] на V [seq_len × dᵥ] даёт выход [seq_len × dᵥ] — взвешенная сумма информации от всех токенов.'),
    ]

    for sym, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{sym}: ')
        run.bold = True
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        p.add_run(desc)

    doc.add_picture(fig_attention_qkv(), width=Inches(5.5))
    doc.add_page_break()

    doc.add_heading('2.4. Пример: что видит модель в предложении о рыбалке', level=2)

    p = doc.add_paragraph()
    p.add_run('Рассмотрим предложение: "Щука поймана на спиннинг в июне".')
    p = doc.add_paragraph()
    p.add_run('Матрица внимания — это таблица 6×6, где на пересечении строки и столбца '
              'стоит "сила внимания" от слова-строки к слову-столбцу.')

    p = doc.add_paragraph()
    p.add_run('Например, когда модель обрабатывает слово "поймана", она:')
    p = doc.add_paragraph('• Сильно "смотрит" на "Щука" (кто пойман?)', style='List Bullet')
    p = doc.add_paragraph('• Умеренно смотрит на "спиннинг" (чем пойман?)', style='List Bullet')
    p = doc.add_paragraph('• Слабо смотрит на "в" (предлог-мусор)', style='List Bullet')

    doc.add_picture(fig_attention_weights(), width=Inches(4.5))
    doc.add_page_break()

    doc.add_heading('2.5. Маскированное внимание (Causal / Masked Attention)', level=2)

    p = doc.add_paragraph()
    p.add_run('Когда мы генерируем текст (не обучаемся), модель не должна "подглядывать" '
              'в будущие слова. Если мы предсказываем 5-е слово, мы можем смотреть только на '
              'слова 1-4.')

    p = doc.add_paragraph()
    p.add_run('Для этого используется ')
    run = p.add_run('маскированное внимание')
    run.bold = True
    p.add_run(': верхний треугольник матрицы внимания (где j > i) зануляется.')

    p = doc.add_paragraph()
    run = p.add_run('Маска = [[1, 0, 0, ...], [1, 1, 0, ...], [1, 1, 1, ...], ...]')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('Без этой маски модель "жульничала" бы — подсматривала ответ, который '
              'должна предсказать. В оригинальном Transformer (машинный перевод) маска '
              'нужна только в декодере. В GPT / LLaMA / нашей модели — всегда.')

    doc.add_heading('2.6. Почему dot-product, а не additive внимание?', level=2)

    p = doc.add_paragraph()
    p.add_run('Vaswani et al. показали, что dot-product attention:')
    p = doc.add_paragraph('В 2-3 раза быстрее (можно использовать оптимизированные матричные умножения)', style='List Bullet')
    p = doc.add_paragraph('Занимает меньше памяти (не надо хранить промежуточные векторы)', style='List Bullet')
    p = doc.add_paragraph('Даёт сопоставимое или лучшее качество', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Это классический trade-off: additive внимание теоретически более выразительно '
              '(может моделировать более сложные функции), но на практике dot-product '
              'выигрывает за счёт эффективности. Для нашей задачи (150M параметров, '
              'обучение на Colab) dot-product — единственный реалистичный выбор.')
    doc.add_page_break()

    # ═══ 3. MULTI-HEAD ═══
    doc.add_heading('Глава 3. Multi-Head Attention — внимание с разных точек зрения', level=1)

    p = doc.add_paragraph()
    p.add_run('Одна матрица внимания — это один "взгляд" на текст. Но в предложении '
              '"Щука поймана на спиннинг в июне" мы хотим одновременно видеть:')

    p = doc.add_paragraph('• Синтаксические связи: подлежащее-сказуемое (Щука — поймана)', style='List Bullet')
    p = doc.add_paragraph('• Семантические связи: орудие-действие (спиннинг — поймана)', style='List Bullet')
    p = doc.add_paragraph('• Временные связи: месяц-событие (июнь — поймана)', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Multi-Head Attention решает эту задачу: делаем h параллельных "голов" внимания, '
              'каждая со своими Q, K, V. Обучаясь, головы специализируются на разных типах связей.')

    p = doc.add_paragraph()
    run = p.add_run('MultiHead(Q, K, V) = Concat(head₁, ..., headₕ) · Wᵒ')
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('где headᵢ = Attention(Q · Wᵢ^Q, K · Wᵢ^K, V · Wᵢ^V). '
              'Каждая голова имеет свои обучаемые проекции Wᵢ^Q, Wᵢ^K, Wᵢ^V.')

    p = doc.add_paragraph()
    p.add_run('Сколько голов? ')
    run = p.add_run('В оригинальном Transformer (2017): h=8 (d_model=512, dₖ=64). '
                     'В GPT-3: h=96 (d_model=12288, dₖ=128). Для нашей модели (150M): '
                     'h=8–12 (d_model=768, dₖ=64).')
    run.bold = True

    doc.add_picture(fig_multihead(), width=Inches(5.5))
    doc.add_page_break()

    # ═══ 4. POSITIONAL ENCODING ═══
    doc.add_heading('Глава 4. Позиционное кодирование — как модель узнаёт порядок слов', level=1)

    p = doc.add_paragraph()
    p.add_run('Ключевое отличие Transformer от RNN: он обрабатывает все слова параллельно, '
              'но теряет информацию о порядке слов. Для модели "щука поймала рыбака" и '
              '"рыбака поймала щука" — одинаковый набор токенов.')

    p = doc.add_paragraph()
    run = p.add_run('Решение: добавить к каждому эмбеддингу сигнал, кодирующий позицию слова. ')
    run.bold = True

    doc.add_heading('4.1. Синусоидальное кодирование (Vaswani et al., 2017)', level=2)

    p = doc.add_paragraph()
    run = p.add_run('PE(pos, 2i) = sin(pos / 10000^{2i/d_model})')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run('PE(pos, 2i+1) = cos(pos / 10000^{2i/d_model})')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run('Где:')
    p = doc.add_paragraph('• pos — позиция слова в предложении (0, 1, 2, ...)', style='List Bullet')
    p = doc.add_paragraph('• i — индекс размерности (0, 1, 2, ..., d_model/2)', style='List Bullet')
    p = doc.add_paragraph('• d_model — размерность модели (например, 768)', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Каждая размерность позиционного кодирования — это синусоида со своей частотой. '
              'Разные размерности "видят" разные масштабы:')

    p = doc.add_paragraph('• Низкие частоты (i малы) — меняются быстро, кодируют близкие позиции')
    p = doc.add_paragraph('• Высокие частоты (i велики) — меняются медленно, кодируют глобальное положение')

    p = doc.add_paragraph()
    p.add_run('Благодаря этому:')
    p = doc.add_paragraph('• Модель может определить абсолютную позицию слова', style='List Bullet')
    p = doc.add_paragraph('• Может понять относительное расстояние между словами', style='List Bullet')
    p = doc.add_paragraph('• Может экстраполировать на предложения длиннее тех, что видела на обучении', style='List Bullet')

    doc.add_picture(fig_positional(), width=Inches(5.5))

    p = doc.add_paragraph()
    p.add_run('Важно: ')
    p.add_run('в современных моделях (GPT-NeoX, LLaMA, Qwen) синусоидальное кодирование '
              'вытеснено RoPE (Rotary Position Embedding). RoPE "поворачивает" Q и K в зависимости '
              'от позиции, что эффективнее и лучше обобщается. Мы будем использовать RoPE в нашей модели.')
    doc.add_page_break()

    # ═══ 5. FULL LAYER ═══
    doc.add_heading('Глава 5. Полный слой декодера — сборка всех компонентов', level=1)

    p = doc.add_paragraph()
    p.add_run('Теперь, когда мы понимаем каждый компонент, соберём их вместе. '
              'Один слой Transformer-декодера состоит из:')

    p = doc.add_paragraph('1. Вход: эмбеддинги токенов + позиционное кодирование')
    p = doc.add_paragraph('2. Layer Normalization — стабилизация распределений активаций')
    p = doc.add_paragraph('3. Multi-Head Self-Attention (маскированное) — обмен информацией между токенами')
    p = doc.add_paragraph('4. + Residual connection — обход внимания', style='List Number')
    p = doc.add_paragraph('5. Layer Normalization — вторая нормализация')
    p = doc.add_paragraph('6. Feed-Forward Network — каждый токен "думает" самостоятельно')
    p = doc.add_paragraph('7. + Residual connection — обход FFN', style='List Number')

    doc.add_picture(fig_decoder_layer(), width=Inches(4.5))

    doc.add_heading('5.1. Feed-Forward Network (FFN)', level=2)

    p = doc.add_paragraph()
    p.add_run('После внимания (которое собрало информацию отовсюду) каждый токен проходит '
              'через полносвязную сеть:')

    p = doc.add_paragraph()
    run = p.add_run('FFN(x) = ReLU(x · W₁ + b₁) · W₂ + b₂')
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('Размер внутреннего слоя обычно в 4 раза больше d_model (например, '
              'd_model=768, d_ff=3072). Это самая "дорогая" по параметрам часть — '
              'примерно 2/3 всех параметров модели находятся в FFN.')

    p = doc.add_paragraph()
    p.add_run('Зачем нужен FFN? ')
    p.add_run('Внимание — это "собрание": токены обмениваются информацией. '
              'FFN — это "кабинет для размышлений": каждый токен обрабатывает '
              'полученную информацию.')

    doc.add_heading('5.2. Residual Connection (остаточная связь)', level=2)

    p = doc.add_paragraph()
    p.add_run('Формула: ')
    run = p.add_run('x_out = x_in + F(x_in)')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run('Зачем? При обучении глубоких сетей (32+ слоя) градиент может затухать. '
              'Residual connection даёт "короткий путь" для градиента — он может "обойти" '
              'слой, если тот не дал прироста. Впервые предложено в ResNet (He et al., 2015) '
              'для компьютерного зрения, затем адаптировано в Transformer.')

    doc.add_heading('5.3. Layer Normalization', level=2)

    p = doc.add_paragraph()
    p.add_run('Вычисляет среднее и дисперсию активаций по скрытым размерностям:')

    p = doc.add_paragraph()
    run = p.add_run('LN(x) = (x - μ) / √(σ² + ε) · γ + β')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run('Зачем? Стабилизирует обучение: активации разных слоёв не "взрываются" '
              'и не "затухают", что позволяет использовать более высокий learning rate.')

    doc.add_heading('5.4. Сколько всего слоёв?', level=2)

    p = doc.add_paragraph()
    p.add_run('Количество слоёв (L) зависит от размера модели:')

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Модель', 'Параметры', 'Слоёв (L)', 'd_model']
    models = [
        ['GPT-2 Small', '124M', '12', '768'],
        ['GPT-2 XL', '1.5B', '48', '1600'],
        ['GPT-3', '175B', '96', '12288'],
        ['Наша модель', '~150M', '12', '768'],
    ]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        for p_run in table.cell(0, j).paragraphs:
            for r in p_run.runs:
                r.bold = True
    for i, row in enumerate(models):
        for j, val in enumerate(row):
            table.cell(i+1, j).text = val
            if i == len(models) - 1:
                for p_run in table.cell(i+1, j).paragraphs:
                    for r in p_run.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

    doc.add_page_break()

    # ═══ 6. TRAINING ═══
    doc.add_heading('Глава 6. Обучение и генерация', level=1)

    doc.add_heading('6.1. Языковое моделирование (next token prediction)', level=2)

    p = doc.add_paragraph()
    p.add_run('LLM обучается предсказывать следующий токен. Для каждого токена в '
              'последовательности модель вычисляет вероятность каждого слова из словаря:')

    p = doc.add_paragraph()
    run = p.add_run('P(xₜ | x₁, x₂, ..., xₜ₋₁) = softmax(hₜ · W_out)')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run('Затем мы сравниваем предсказание с правильным ответом через '
              'кросс-энтропию (cross-entropy loss):')

    p = doc.add_paragraph()
    run = p.add_run('L = -Σ log P(правильный_токен_t | x_{<t})')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run('Чем меньше loss, тем увереннее модель предсказывает правильный токен.')

    doc.add_heading('6.2. Авторегрессивная генерация', level=2)

    p = doc.add_paragraph()
    p.add_run('На этапе генерации мы:')
    p = doc.add_paragraph('Подаём начальную строку (промпт)', style='List Number')
    p = doc.add_paragraph('Модель вычисляет вероятности для следующего токена', style='List Number')
    p = doc.add_paragraph('Выбираем токен (разные стратегии — см. ниже)', style='List Number')
    p = doc.add_paragraph('Добавляем его к контексту', style='List Number')
    p = doc.add_paragraph('Повторяем до токена <endoftext> или макс. длины', style='List Number')

    doc.add_picture(fig_generation(), width=Inches(5.5))

    doc.add_heading('6.3. Стратегии выбора токена', level=2)

    p = doc.add_paragraph()
    p.add_run('Greedy (жадный): ')
    run = p.add_run('всегда выбираем токен с наибольшей вероятностью. '
                     'Просто, но предсказуемо, склонно к повторам.')
    run.font.italic = True

    p = doc.add_paragraph()
    p.add_run('Top-k: ')
    run = p.add_run('выбираем из k самых вероятных токенов (k=40-50). '
                     'Больше разнообразия.')
    run.font.italic = True

    p = doc.add_paragraph()
    p.add_run('Top-p (nucleus): ')
    run = p.add_run('выбираем из минимального набора токенов, чья суммарная '
                     'вероятность превышает p (p=0.9). Адаптивно.')
    run.font.italic = True

    p = doc.add_paragraph()
    p.add_run('Temperature: ')
    run = p.add_run('масштабирует распределение. T=0.1 → почти greedy. '
                     'T=1.0 → оригинальное распределение. T=2.0 → более хаотично. '
                     'Полезно для творческих ответов.')
    run.font.italic = True

    p = doc.add_paragraph()
    run = p.add_run('Для нашей рыболовной модели: ')
    p.add_run('top-p=0.9 + temperature=0.7 — хороший баланс '
              'между точностью и разнообразием.')
    doc.add_page_break()

    # ═══ 7. WHY TRANSFORMER WON ═══
    doc.add_heading('Глава 7. Почему Transformer "победил" все альтернативы', level=1)

    p = doc.add_paragraph()
    p.add_run('1. Параллелизация: ')
    p.add_run('все токены обрабатываются за O(1) шагов (не O(n) как в RNN). '
              'Это позволило обучать на огромных корпусах.')

    p = doc.add_paragraph()
    p.add_run('2. Дальние зависимости: ')
    p.add_run('любой токен "видит" любой другой за один шаг (а не через n шагов). '
              'Решена проблема затухающих градиентов.')

    p = doc.add_paragraph()
    p.add_run('3. Масштабируемость: ')
    p.add_run('Transformer показал "scaling laws" — качество предсказуемо растёт '
              'с увеличением параметров, данных и compute (Kaplan et al., 2020).')

    p = doc.add_paragraph()
    p.add_run('4. Индуктивные bias: ')
    p.add_run('Transformer не навязывает сильных предубеждений (как локальность в CNN '
              'или последовательность в RNN), что позволяет ему учить любые паттерны.')

    p = doc.add_paragraph()
    p.add_run('5. Эмбеддинги: ')
    p.add_run('скрытые представления токенов можно использовать для transfer learning '
              '(fine-tuning). Это то, что мы и будем делать — модель, обученная на общих данных, '
              'дообучается на нашем рыбацком корпусе.')

    p = doc.add_paragraph()
    run = p.add_run('Недостатки Transformer: ')
    p.add_run('главный — квадратичная сложность O(n²) по длине последовательности '
              '(матрица внимания n×n). Для нашей модели (контекст 1024-2048 токенов) '
              'это приемлемо. Для 100K+ токенов нужны оптимизации (Flash Attention, '
              'Sparse Attention). Мы будем использовать Flash Attention.')
    doc.add_page_break()

    # ═══ 8. EXERCISES ═══
    doc.add_heading('Глава 8. Практические задания', level=1)

    p = doc.add_paragraph()
    p.add_run('Для закрепления материала выполните следующие задания:')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Задание 1. Минимальное внимание на NumPy')
    p.runs[0].bold = True
    code = doc.add_paragraph()
    run = code.add_run(
        '# Реализуйте self-attention для предложения из 3 слов:\n'
        '# "щука поймана на"\n'
        '# Размерность: d_k = 4\n\n'
        'import numpy as np\n\n'
        'def softmax(x, axis=-1):\n'
        '    e_x = np.exp(x - np.max(x, axis=axis, keepdims=True))\n'
        '    return e_x / e_x.sum(axis=axis, keepdims=True)\n\n'
        '# Дано: X = [[1, 0, 1, 0],\n'
        '#            [0, 1, 0, 1],\n'
        '#            [1, 1, 0, 0]]\n'
        'X = np.array([[1, 0, 1, 0],\n'
        '              [0, 1, 0, 1],\n'
        '              [1, 1, 0, 0]])\n\n'
        'Q = K = V = X  # для простоты\n\n'
        'scores = Q @ K.T / np.sqrt(4)\n'
        'weights = softmax(scores)\n'
        'output = weights @ V\n\n'
        'print("Матрица внимания:")\n'
        'print(weights)\n'
        'print("\\nВыход:")\n'
        'print(output)\n\n'
        '# Вопрос: какое слово имеет самую "сильную" связь с "щука"?'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Задание 2. Анализ реальной модели')
    p.runs[0].bold = True
    p = doc.add_paragraph()
    p.add_run('Откройте модель LLaMA или GPT-2 на Hugging Face Model Hub. '
              'Найдите в конфиге (config.json):')
    p = doc.add_paragraph('• d_model', style='List Bullet')
    p = doc.add_paragraph('• n_heads', style='List Bullet')
    p = doc.add_paragraph('• n_layers', style='List Bullet')
    p = doc.add_paragraph('• vocab_size', style='List Bullet')
    p = doc.add_paragraph('• d_ff (или intermediate_size)', style='List Bullet')
    p = doc.add_paragraph('Посчитайте общее количество параметров: '
                          'N = vocab_size·d_model + L·(4·d_model² + 8·d_model·d_k·n_heads) + d_model·vocab_size')

    p = doc.add_paragraph()
    p.add_run('Задание 3. Визуализация внимания')
    p.runs[0].bold = True
    p = doc.add_paragraph()
    p.add_run('Напишите функцию, которая вычисляет и визуализирует матрицу внимания '
              'для предложения "Лещ клюёт на червя в июне на Верхней Волге". '
              'Какая пара слов имеет максимальное значение внимания? Почему?')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Задание 4 (со звёздочкой). Сравнение с RNN')
    p.runs[0].bold = True
    p = doc.add_paragraph()
    p.add_run('Реализуйте простую RNN с одним скрытым слоем (d_hidden=4) и '
              'обучите её на синтетических данных — последовательностях, где важно '
              'помнить слово из начала. Сравните с однослойным Transformer. '
              'Какая архитектура лучше справляется при длине последовательности 20? 50?')
    doc.add_page_break()

    # ═══ 9. SUMMARY ═══
    doc.add_heading('Резюме: что нужно вынести из этой лекции', level=1)

    items = [
        'Transformer работает параллельно (не последовательно, как RNN) — это ключ к его эффективности',
        'Self-Attention — механизм, позволяющий каждому токену "смотреть" на все остальные. Q — запрос, K — индекс, V — содержание',
        'Multi-Head Attention — несколько параллельных "взглядов" на текст (8-96 голов)',
        'Positional Encoding — способ сообщить модели порядок слов. В нашей модели будем использовать RoPE',
        'Residual connection — обходной путь, позволяющий обучать глубокие сети (12-96 слоёв)',
        'FFN — "размышление" каждого токена после "совещания" (Attention). Занимает ~2/3 параметров модели',
        'LLM обучается на предсказание следующего токена (next token prediction)',
        'На этапе генерации используем top-p=0.9 + temperature=0.7 для хорошего баланса',
        'Transformer имеет недостатки (O(n²) сложность), но для нашей задачи (контекст ~2K токенов) это не проблема',
        'Современные LLM = Transformer Decoder + RoPE + (SiLU/GeLU) + QK Normalization',
    ]
    for i, item in enumerate(items):
        doc.add_paragraph(f'{i+1}. {item}')

    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('Литература и источники:')
    run.bold = True
    run.font.size = Pt(12)

    refs = [
        'Vaswani et al. (2017) — "Attention Is All You Need". https://arxiv.org/abs/1706.03762',
        'Hochreiter & Schmidhuber (1997) — "Long Short-Term Memory". Neural Computation.',
        'Bahdanau et al. (2015) — "Neural Machine Translation by Jointly Learning to Align and Translate". https://arxiv.org/abs/1409.0473',
        'He et al. (2015) — "Deep Residual Learning for Image Recognition". https://arxiv.org/abs/1512.03385',
        '3Blue1Brown (2024) — "Attention in transformers, visually explained". YouTube.',
        'Karpathy (2023) — "Let\'s build GPT from scratch". YouTube / GitHub (nanoGPT).',
        'Kaplan et al. (2020) — "Scaling Laws for Neural Language Models". https://arxiv.org/abs/2001.08361',
        'Wolf et al. (2020) — "Transformers: State-of-the-Art Natural Language Processing". https://arxiv.org/abs/1910.03771',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run(
        'Следующая лекция: Лекция 2 — Tokenization: как текст превращается в числа. '
        'Будем разбирать BPE, SentencePiece и обучим свой токенизатор на корпусе рыбацких текстов.'
    )
    run.bold = True
    run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, 'lecture-01-transformer.docx')
    doc.save(path)
    print(f'Документ сохранён: {path}')


if __name__ == '__main__':
    build_doc()
