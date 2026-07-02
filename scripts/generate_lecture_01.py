"""Генерация Word-документа с иллюстрациями: Лекция 1 — Transformer"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import io
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')

# === Цветовая палитра ===
COLORS = {
    'bg': '#1a1a2e',
    'primary': '#16213e',
    'secondary': '#0f3460',
    'accent': '#e94560',
    'orange': '#f59e0b',
    'green': '#10b981',
    'blue': '#3b82f6',
    'purple': '#8b5cf6',
    'teal': '#06b6d4',
    'text': '#e2e8f0',
    'text_dim': '#94a3b8',
}


def setup_style(ax):
    """Единый стиль для графиков"""
    ax.set_facecolor(COLORS['bg'])
    ax.figure.patch.set_facecolor(COLORS['bg'])
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color(COLORS['text_dim'])
    ax.spines['bottom'].set_color(COLORS['text_dim'])
    ax.tick_params(colors=COLORS['text_dim'])
    ax.grid(True, alpha=0.1, color=COLORS['text_dim'])


def save_fig(fig):
    """Сохраняем фигуру в BytesIO для вставки в Word"""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor=COLORS['bg'], edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


# ============================================================
# ДИАГРАММА 1: Высокоуровневая схема LLM (Decoder-only)
# ============================================================
def diagram_overview():
    fig, ax = plt.subplots(figsize=(8, 7))
    setup_style(ax)
    ax.set_xlim(-1, 1)
    ax.set_ylim(-1, 1)
    ax.axis('off')

    boxes = [
        (-0.3, 0.85, 0.6, 0.10, 'OUTPUT\n("щука", "12", "кг")', COLORS['accent']),
        (-0.4, 0.65, 0.8, 0.10, 'DECODER (LLM)', COLORS['secondary']),
        (-0.4, 0.45, 0.8, 0.10, 'Tokenizer → Embedding', COLORS['purple']),
        (-0.5, 0.20, 1.0, 0.12, 'INPUT: "Щука весила 12 кг"', COLORS['green']),
    ]

    for x, y, w, h, label, color in boxes:
        rect = FancyBboxPatch((x, y), w, h,
                              boxstyle="round,pad=0.02",
                              facecolor=color, alpha=0.9,
                              edgecolor='white', linewidth=0.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, label,
                ha='center', va='center', fontsize=9, color='white',
                fontweight='bold')

    # Стрелки
    for y_from, y_to in [(0.76, 0.86), (0.56, 0.66), (0.33, 0.46)]:
        ax.annotate('', xy=(0, y_from), xytext=(0, y_to),
                    arrowprops=dict(arrowstyle='->', color=COLORS['text'],
                                    lw=2))

    # Комментарий сбоку
    ax.text(0.55, 0.55, 'Decoder-Only\n(GPT-style)',
            fontsize=8, color=COLORS['orange'], fontweight='bold')
    ax.text(0.55, 0.25, 'Токенизация +\nEmbedding',
            fontsize=8, color=COLORS['teal'], fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. 1: Высокоуровневая архитектура языковой модели',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 2: Токенизация (BPE)
# ============================================================
def diagram_tokenization():
    fig, ax = plt.subplots(figsize=(8, 3.5))
    setup_style(ax)
    ax.set_xlim(-1, 6)
    ax.set_ylim(0, 1)
    ax.axis('off')

    colors_token = [COLORS['blue'], COLORS['purple'], COLORS['teal'],
                    COLORS['green'], COLORS['orange'], COLORS['accent']]

    sentence = "Щука весила 12 кг"
    tokens = [("Щук", 0.0), ("а", 1.0), ("_", 2.0),
              ("вес", 3.0), ("ила", 4.0), ("_", 5.0),
              ("12", 6.0), ("_", 7.0), ("кг", 8.0)]

    for i, (tok, x) in enumerate(tokens):
        color = colors_token[i % len(colors_token)]
        rect = FancyBboxPatch((x - 0.35, 0.3), 0.7, 0.4,
                              boxstyle="round,pad=0.05",
                              facecolor=color, alpha=0.85,
                              edgecolor='white', linewidth=0.5)
        ax.add_patch(rect)
        ax.text(x, 0.5, tok, ha='center', va='center', fontsize=10,
                color='white', fontweight='bold')

    # Стрелка от исходного предложения
    ax.text(4, 0.85, 'Исходный текст: "Щука весила 12 кг"',
            ha='center', fontsize=11, color=COLORS['text'], fontweight='bold')

    ax.annotate('', xy=(4, 0.73), xytext=(4, 0.85),
                arrowprops=dict(arrowstyle='->', color=COLORS['text'], lw=2))

    # Стрелки вниз — куда идут токены
    ax.text(4, 0.15, '→ Токен ID: [481, 203, 307, 1562, 890, 307, 512, 307, 901]',
            ha='center', fontsize=9, color=COLORS['orange'])
    ax.text(4, 0.02, '→ Embedding Table (32K × 4096)',
            ha='center', fontsize=9, color=COLORS['teal'])

    fig.text(0.5, 0.01, 'Рис. 2: Byte Pair Encoding — разбиение текста на токены',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 3: Self-Attention QKV
# ============================================================
def diagram_attention():
    fig, ax = plt.subplots(figsize=(8, 5))
    setup_style(ax)
    ax.set_xlim(-1, 5)
    ax.set_ylim(0, 3.5)
    ax.axis('off')

    # Три колонки: Q, K, V с матрицами
    for col_idx, (name, x_off, color) in enumerate([
        ('Q (Query)', 0.2, COLORS['blue']),
        ('K (Key)', 1.8, COLORS['purple']),
        ('V (Value)', 3.4, COLORS['green']),
    ]):
        # Заголовок
        ax.text(x_off + 0.4, 3.0, name, fontsize=10, color=color,
                fontweight='bold', ha='center')

        # Матрица
        for row in range(4):
            for col in range(3):
                val = np.random.randint(0, 10)
                alpha = 0.3 + 0.6 * (val / 10)
                rect = FancyBboxPatch(
                    (x_off + col * 0.3, 1.5 - row * 0.3),
                    0.28, 0.28, boxstyle="round,pad=0.01",
                    facecolor=color, alpha=alpha,
                    edgecolor='white', linewidth=0.3)
                ax.add_patch(rect)
                ax.text(x_off + col * 0.3 + 0.14, 1.5 - row * 0.3 + 0.14,
                        str(val), ha='center', va='center',
                        fontsize=6, color='white')

        # Текст снизу
        info = {
            'Q (Query)': '"Что я ищу?"\nКаждое слово формирует\nвопрос к другим словам',
            'K (Key)': '"Что у меня есть?"\nКаждое слово отвечает,\nнасколько оно релевантно',
            'V (Value)': '"Что я передаю?"\nИнформация, которую\nслово отдает другим',
        }[name]
        ax.text(x_off + 0.4, 0.5, info, fontsize=7, color=COLORS['text'],
                ha='center', va='center')

    # Формула
    ax.text(2.5, 3.2, 'Attention(Q,K,V) = softmax(Q·Kᵀ / √d) · V',
            ha='center', fontsize=9, color=COLORS['accent'],
            fontweight='bold', family='monospace')

    fig.text(0.5, 0.02, 'Рис. 3: Механизм Self-Attention: Query, Key, Value',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 4: Multi-Head Attention
# ============================================================
def diagram_multihead():
    fig, ax = plt.subplots(figsize=(7, 5))
    setup_style(ax)
    ax.set_xlim(-1, 1)
    ax.set_ylim(-1, 1)
    ax.axis('off')

    # Вход
    rect = FancyBboxPatch((-0.3, 0.75), 0.6, 0.1,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['green'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, 0.8, 'INPUT', ha='center', va='center',
            fontsize=9, color='white', fontweight='bold')

    # Разделение на головы
    y = 0.5
    for i in range(4):
        x = -0.35 + i * 0.22
        rect = FancyBboxPatch((x, y), 0.18, 0.15,
                              boxstyle="round,pad=0.02",
                              facecolor=COLORS['purple'], alpha=0.6 + i*0.1,
                              edgecolor='white', linewidth=0.5)
        ax.add_patch(rect)
        ax.text(x + 0.09, y + 0.075, f'Head {i+1}',
                ha='center', va='center', fontsize=7, color='white')

    # Стрелки вниз от каждого внимания
    y2 = 0.2
    for i in range(4):
        x = -0.35 + i * 0.22
        ax.annotate('', xy=(x + 0.09, y2 + 0.1), xytext=(x + 0.09, y + 0.15),
                    arrowprops=dict(arrowstyle='->', color=COLORS['text'], lw=1,
                                    alpha=0.5))

    # Concat
    rect = FancyBboxPatch((-0.4, 0.0), 0.8, 0.1,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['orange'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, 0.05, 'CONCATENATE (склейка)', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Linear projection
    rect = FancyBboxPatch((-0.3, -0.25), 0.6, 0.1,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['accent'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, -0.2, 'LINEAR PROJECTION', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Стрелка от Concat к Linear
    ax.annotate('', xy=(0, -0.25), xytext=(0, 0),
                arrowprops=dict(arrowstyle='->', color=COLORS['text'], lw=2))

    # Подписи
    ax.text(0.55, 0.5, 'Каждая голова\nищет свои\nпаттерны:\n\nГолова 1: объект-действие\nГолова 2:\nприл-сущ\nГолова 3:\nчисла\n...', fontsize=7,
            color=COLORS['text_dim'], verticalalignment='center')

    fig.text(0.5, 0.02, 'Рис. 4: Multi-Head Attention — несколько "взглядов" на текст',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 5: Один слой Transformer
# ============================================================
def diagram_decoder_layer():
    fig, ax = plt.subplots(figsize=(6, 7))
    setup_style(ax)
    ax.set_xlim(-1, 1)
    ax.set_ylim(-1, 1)
    ax.axis('off')

    # Вход
    rect = FancyBboxPatch((-0.35, 0.80), 0.7, 0.08,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['green'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, 0.84, 'Input Tokens', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Layer Norm 1
    rect = FancyBboxPatch((-0.30, 0.62), 0.6, 0.08,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['teal'], alpha=0.8,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, 0.66, 'Layer Norm', ha='center', va='center',
            fontsize=8, color='white')

    # Multi-Head Attention
    rect = FancyBboxPatch((-0.35, 0.43), 0.7, 0.12,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['purple'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, 0.49, 'Multi-Head Attention', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Residual connection 1 (пунктиром)
    ax.annotate('', xy=(0.38, 0.80), xytext=(0.38, 0.44),
                arrowprops=dict(arrowstyle='->', color=COLORS['orange'],
                                lw=1.5, linestyle='dashed'))
    ax.text(0.42, 0.62, '+', fontsize=12, color=COLORS['orange'],
            fontweight='bold')

    # Layer Norm 2
    rect = FancyBboxPatch((-0.30, 0.27), 0.6, 0.08,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['teal'], alpha=0.8,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, 0.31, 'Layer Norm', ha='center', va='center',
            fontsize=8, color='white')

    # FFN
    rect = FancyBboxPatch((-0.35, 0.08), 0.7, 0.12,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['blue'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, 0.14, 'Feed-Forward Network', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Residual connection 2
    ax.annotate('', xy=(0.38, 0.43), xytext=(0.38, 0.10),
                arrowprops=dict(arrowstyle='->', color=COLORS['orange'],
                                lw=1.5, linestyle='dashed'))
    ax.text(0.42, 0.26, '+', fontsize=12, color=COLORS['orange'],
            fontweight='bold')

    # Выход
    rect = FancyBboxPatch((-0.25, -0.08), 0.5, 0.08,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['accent'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(0, -0.04, 'OUTPUT → Next Layer', ha='center', va='center',
            fontsize=8, color='white', fontweight='bold')

    # Стрелки
    y_positions = [(0.80, 0.63), (0.62, 0.45), (0.44, 0.28), (0.27, 0.10)]
    for y1, y2 in y_positions:
        ax.annotate('', xy=(0, y1), xytext=(0, y2),
                    arrowprops=dict(arrowstyle='->', color=COLORS['text'], lw=1.5))

    # Легенда
    ax.text(-0.9, 0.9, 'Residual\nconnection\n(обходной\nпуть)',
            fontsize=7, color=COLORS['orange'], fontweight='bold')

    fig.text(0.5, 0.02, 'Рис. 5: Один слой Decoder-Transformer',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# ДИАГРАММА 6: Процесс генерации текста
# ============================================================
def diagram_generation():
    fig, ax = plt.subplots(figsize=(9, 4))
    setup_style(ax)
    ax.set_xlim(-1, 10)
    ax.set_ylim(0, 1)
    ax.axis('off')

    # Шаг 1: Исходный вход
    input_text = "Какая наживка нужна для леща"
    rect = FancyBboxPatch((0.0, 0.5), 3.5, 0.35,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['green'], alpha=0.85,
                          edgecolor='white', linewidth=0.5)
    ax.add_patch(rect)
    ax.text(1.75, 0.675, f'INPUT: "{input_text}..."',
            ha='center', va='center', fontsize=9, color='white',
            fontweight='bold')

    # LLM блок
    rect = FancyBboxPatch((4.0, 0.45), 2.0, 0.45,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['secondary'], alpha=0.9,
                          edgecolor=COLORS['accent'], linewidth=2)
    ax.add_patch(rect)
    ax.text(5.0, 0.675, 'LLM\n(Transformer)', ha='center', va='center',
            fontsize=10, color='white', fontweight='bold')

    # Распределение вероятностей
    probs = [0.02, 0.01, 0.15, 0.03, 0.02, 0.35, 0.12, 0.05, 0.20, 0.05]
    tokens_out = ['в', 'на', 'лучше', 'после', 'около', 'в\nиюне', 'там',
                  'где', 'на\nВолге', 'это']
    for i, (p, tok) in enumerate(zip(probs, tokens_out)):
        x = 6.5 + i * 0.32
        height = p * 6
        rect = FancyBboxPatch((x - 0.08, 0.05), 0.16, height,
                              boxstyle="round,pad=0.01",
                              facecolor=COLORS['accent'],
                              alpha=0.3 + p * 2,
                              edgecolor='white', linewidth=0.5)
        ax.add_patch(rect)
        ax.text(x, height + 0.08, f'{p:.0%}', ha='center', va='bottom',
                fontsize=5, color=COLORS['text'])
        ax.text(x, 0.0, tok, ha='center', va='top',
                fontsize=5.5, color=COLORS['text_dim'])

    # Highlights
    highlight_idx = 5  # "в июне"
    x_hl = 6.5 + highlight_idx * 0.32
    rect = FancyBboxPatch((x_hl - 0.12, 0.02), 0.24, 0.3,
                          boxstyle="round,pad=0.02",
                          facecolor=COLORS['orange'], alpha=0.3,
                          edgecolor=COLORS['orange'], linewidth=1.5,
                          linestyle='dashed')
    ax.add_patch(rect)
    ax.text(x_hl, 0.38, '✓ выбран\n"в июне"', ha='center', va='bottom',
            fontsize=5, color=COLORS['orange'], fontweight='bold')

    # Стрелка
    ax.annotate('', xy=(4.0, 0.675), xytext=(3.55, 0.675),
                arrowprops=dict(arrowstyle='->', color=COLORS['text'], lw=2))
    ax.annotate('', xy=(6.0, 0.675), xytext=(6.5, 0.675),
                arrowprops=dict(arrowstyle='->', color=COLORS['text'], lw=2))

    fig.text(0.5, 0.02, 'Рис. 6: Генерация — модель предсказывает следующее слово с вероятностью',
             ha='center', fontsize=8, color=COLORS['text_dim'])
    return save_fig(fig)


# ============================================================
# СБОРКА WORD-ДОКУМЕНТА
# ============================================================
def create_docx():
    doc = Document()

    # Стили
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Титул
    title = doc.add_heading('Лекция 1: Transformer — архитектура, на которой держится всё', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph('')
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run('FishingLLM — Создание экспертной системы по рыбалке в Тверской области\n'
                        'Цель: понять, как устроена нейросеть, которая будет нашим "рыболовным экспертом"')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # === 1 ===
    doc.add_heading('1. Проблема, которую решал Transformer (2017)', level=1)
    p = doc.add_paragraph()
    p.add_run('До 2017 года все нейросети для текста были последовательными:\n').bold = True
    p = doc.add_paragraph()
    run = p.add_run('• RNN / LSTM ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
    p.add_run('— читают слово за словом, слева направо. Чем длиннее предложение, тем хуже помнят начало.\n')
    run = p.add_run('• CNN для текста ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
    p.add_run('— видят только локальное окно в 5–7 слов.')

    p = doc.add_paragraph()
    p.add_run('Проблема: ')
    run = p.add_run('В предложении "Щука, которую я поймал вчера на спиннинг возле устья Медведицы,'
                     ' весила 12 кг" — слова "вчера" и "устье Медведицы" влияют на смысл далеко друг от друга.'
                     ' RNN это "забывает".')
    run.font.color.rgb = RGBColor(0xE9, 0x45, 0x60)

    p = doc.add_paragraph()
    p.add_run('Решение: ')
    run = p.add_run('Transformer читает все слова сразу, используя механизм "внимания" (Attention),'
                     ' чтобы модель сама решала, на какие слова смотреть при предсказании.')
    run.bold = True

    doc.add_picture(diagram_overview(), width=Inches(5))
    doc.add_paragraph('')
    doc.add_page_break()

    # === 2 ===
    doc.add_heading('2. Токенизация — кирпичики текста', level=1)
    p = doc.add_paragraph()
    p.add_run('Современные токенизаторы не режут по буквам и не по словам. Они используют '
              'Byte Pair Encoding (BPE):')
    p = doc.add_paragraph()
    p.add_run('Строка "Щука весила 12 кг" → [Щук], [а], [_], [вес], [ила], [_], [12], [_], [кг]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_picture(diagram_tokenization(), width=Inches(5.5))
    p = doc.add_paragraph()
    p.add_run('Почему не по словам? ').bold = True
    p.add_run('Словарный запас языка бесконечен ("подлещик", "уклейка", "чехонь"). '
              'BPE находит частые куски и кодирует их. Редкие слова собирает из подслов.')
    doc.add_page_break()

    # === 3 ===
    doc.add_heading('3. Embedding Layer — перевод в числа', level=1)
    p = doc.add_paragraph()
    p.add_run('Каждый токен — это число из словаря (например, "Щук" = 481, "а" = 203). '
              'Модели нужны векторы, а не индексы.')
    p = doc.add_paragraph()
    p.add_run('Embedding — это большая таблица:\n')
    p.add_run('• Строки: все токены словаря (32K–128K)\n'
              '• Столбцы: размерность (обычно 4096 для 7B моделей)\n'
              '• Значения: обучаемые параметры сети')
    p = doc.add_paragraph()
    run = p.add_run('Каждое слово превращается в направление в многомерном пространстве. '
                    'Похожие слова оказываются рядом — это семантическая близость.')
    run.font.italic = True
    doc.add_page_break()

    # === 4 ===
    doc.add_heading('4. Self-Attention — сердце модели', level=1)
    p = doc.add_paragraph()
    p.add_run('Это самый важный механизм. Он позволяет каждому токену "посмотреть" на все '
              'остальные и решить, какие из них важны для предсказания следующего.')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Аналогия: ').bold = True
    p.add_run('На вопрос "Какую наживку лучше использовать для леща в июне на Верхней Волге?" '
              'ты мысленно смотришь на слова "лещ" (целевая рыба), "июнь" (сезон), '
              '"Верхняя Волга" (локация). Attention делает то же самое — вычисляет "вес" важности для каждой пары слов.')

    doc.add_picture(diagram_attention(), width=Inches(5.5))
    doc.add_paragraph()
    doc.add_page_break()

    # === 5 ===
    doc.add_heading('5. Multi-Head Attention', level=1)
    p = doc.add_paragraph()
    p.add_run('Одной головы внимания мало — как смотреть на мир одним глазом. Делаем 32–64 "головы", '
              'каждая ищет свои паттерны:\n'
              '• Голова 1: ищет отношения глагол-дополнение\n'
              '• Голова 2: ищет связи прилагательное-существительное\n'
              '• Голова 3: ищет числовые величины\n'
              '• ...')
    doc.add_picture(diagram_multihead(), width=Inches(5))
    doc.add_page_break()

    # === 6 ===
    doc.add_heading('6. Полный слой Transformer', level=1)
    p = doc.add_paragraph()
    p.add_run('Собираем всё вместе. Один слой состоит из:\n\n'
              '1. Layer Normalization — стабилизирует обучение\n'
              '2. Multi-Head Attention — обмен информацией между токенами\n'
              '3. Residual connection — "запоминалка", если слой ничего не сделал\n'
              '4. Layer Normalization — ещё одна нормализация\n'
              '5. Feed-Forward Network — каждый токен "думает" самостоятельно\n'
              '6. Residual connection — второй обходной путь')
    doc.add_picture(diagram_decoder_layer(), width=Inches(4.5))
    p = doc.add_paragraph()
    run = p.add_run('В модели 7B параметров таких слоёв — 32. В модели 70B — 80.')
    run.bold = True
    doc.add_page_break()

    # === 7 ===
    doc.add_heading('7. Как LLM генерирует текст', level=1)
    p = doc.add_paragraph()
    p.add_run('На этапе обучения: ')
    p.add_run('"Щука весила" → модель предсказывает → "12"').bold = True
    p = doc.add_paragraph()
    p.add_run('На этапе генерации:')
    p = doc.add_paragraph()
    p.add_run('1. Подаём начальную строку\n'
              '2. Модель предсказывает распределение вероятностей следующего токена\n'
              '3. Выбираем токен (случайно по top-p/top-k)\n'
              '4. Добавляем его к строке\n'
              '5. Повторяем, пока не получим токен "конца строки"')
    doc.add_picture(diagram_generation(), width=Inches(6))
    doc.add_page_break()

    # === Итог ===
    doc.add_heading('8. Что важно запомнить', level=1)
    items = [
        'LLM — это предсказатель следующего слова (ничего больше)',
        'Transformer = Attention + FFN × много слоёв',
        'Attention позволяет смотреть на всё сразу, а не по порядку',
        'Параметры модели — это числа в весах, которые обучаются',
        'Размер модели = сумма всех обучаемых чисел (7B = 7 миллиардов)',
        'Для генерации текста нужен только декодер (GPT, LLaMA, Mistral, Qwen)',
    ]
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f'{i}. {item}', style='List Number')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Следующий шаг: Лекция 2 — Tokenization на практике')
    run.font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
    run.bold = True

    # Сохраняем
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, 'lecture-01-transformer.docx')
    doc.save(path)
    print(f'Документ сохранён: {path}')


if __name__ == '__main__':
    create_docx()
